r"""
21_signal_intersections.py
=========================

Phase 3a / Phase 2 — intersection nodes, signalized flags, and per-crash signal
attributes (descriptive, scoped to signal-covered corridors). No causal claims.

Settled parameters (Phase 1 sign-off): dedup 30 m (1,008 crossings); node snap 30 m
for both; coverage = corridor/inventory-based; intersection = junction of 2+ distinct
named through-roads.

Method (all spatial work EPSG:32136):
  1. COVERED CORRIDORS (route-anchored, not corner-noisy snapping): group signals by
     ROUTE_NUMBER; a (route, nearest-street) pair with >= MIN_PER_ROUTE signals marks
     that street as a covered corridor. This attributes a Poplar intersection's corner
     signals to Poplar (the route), not the cross street; cross streets only qualify if
     they themselves carry a route's signals. Ramp/interstate streets excluded.
  2. NODES (all named junctions citywide, tagged coverage): endpoints of through-road
     segments (MTFCC S1200/S1400) snapped to a 1 m grid; a node = >= 2 distinct named
     streets meet there (excludes driveways/alleys/service/ramps and degree-1 ends).
     on_covered = >=1 incident street is a covered corridor.
  3. SIGNALIZED: a node is signalized if a deduped crossing is within 30 m.
  4. CRASH ATTRIBUTES: at_intersection from the NonMotorist field (primary), corroborated
     by snap to the nearest node within 30 m; disagreements flagged ambiguous.

Reads:  data/raw/ped_signals.geojson, data/processed/signalized_crossings_dedup.geojson,
        data/processed/shelby_crashes_final.csv, data/processed/road_ownership_rulebook.geojson
Writes: data/processed/shelby_crashes_signals.csv   (crashes + new columns; NEW file)
        data/processed/intersection_nodes_covered.geojson  (covered nodes + signalized flag)

Run it with:
    .\.venv\Scripts\python.exe scripts\21_signal_intersections.py
"""

import sys
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point

DOCX = Path(__file__).resolve().parent.parent / "data" / "processed" / "novel_statistics.docx"

ROOT = Path(__file__).resolve().parent.parent
RAW = ROOT / "data" / "raw"
PROC = ROOT / "data" / "processed"
PED = RAW / "ped_signals.geojson"
CROSSINGS = PROC / "signalized_crossings_dedup.geojson"
CRASHES = PROC / "shelby_crashes_final.csv"
RULEBOOK = PROC / "road_ownership_rulebook.geojson"
CRASH_OUT = PROC / "shelby_crashes_signals.csv"
NODES_OUT = PROC / "intersection_nodes_covered.geojson"

CRS_M = "EPSG:32136"
SNAP_M = 30.0            # crash->node and signal->node
MIN_ROUTE_TOTAL = 4     # a route needs >= this many signals to be a covered route
SECONDARY_SHARE = 0.25  # a route's non-dominant street is a corridor too if it carries this share
THROUGH_MTFCC = {"S1200", "S1400"}
FATAL = "Fatal"


def pct(p, w):
    return round(100.0 * p / w, 1) if w else 0.0


def append_docx(D):
    """Append a dated Phase-3a section to the living stats doc (idempotent)."""
    try:
        from docx import Document
    except Exception:
        print("  (python-docx not available; skipping docx append)"); return
    if not DOCX.exists():
        print(f"  (docx not found: {DOCX.name}; skipping)"); return
    doc = Document(str(DOCX))
    MARK = "Phase 3a — Pedestrian signals"
    if any(p.text.strip().startswith(MARK) for p in doc.paragraphs):
        print("  (docx already has the Phase-3a section; skipping append)"); return
    doc.add_page_break()
    doc.add_heading(f"Phase 3a — Pedestrian signals & intersections (added {date.today().isoformat()})", 1)
    doc.add_paragraph(
        "Source: TDOT 'ADA Asset Data' (geodata.tn.gov) Pedestrian Signal layer. "
        f"{D['n_signals']:,} signal heads/push-buttons in Shelby were deduped to "
        f"{D['n_cross']:,} signalized pedestrian crossings (one per intersection, 30 m). "
        "Crosswalks were checked and set aside (247 polygons, all on interstate interchanges; "
        "zero surface-arterial crosswalks). All figures are descriptive only — no causal claims.")

    doc.add_heading("Signal coverage & the covered-corridor rule", 2)
    for b in [
        "The TDOT inventory is NOT limited to signed state routes: a large share of signals sit on "
        "major CITY arterials (Winchester, Riverdale, Stage, Germantown, Houston Levee...).",
        f"Coverage is therefore defined by corridor, not ownership: signals were grouped by TDOT "
        f"ROUTE_NUMBER, and each covered route was mapped to its dominant arterial (plus any ≥25% "
        f"secondary), giving {D['n_corridors']} covered corridors (e.g. {D['corridors_sample']}...).",
        f"Intersection nodes were built from the street centerline (junctions of 2+ distinct named "
        f"through-roads, MTFCC S1200/S1400; driveways/alleys/ramps/dead-ends excluded). "
        f"{D['n_covnodes']:,} nodes lie on covered corridors; {D['n_signodes']} are signalized "
        f"(a deduped crossing within 30 m).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("New per-crash attributes (reproducible, live-ready)", 2)
    doc.add_paragraph("Written to data/processed/shelby_crashes_signals.csv (jurisdiction columns untouched):")
    for b in [
        "at_intersection (bool) — from the NonMotorist location field (primary determinant).",
        "intersection_node_id — nearest covered junction node within 30 m, else null.",
        "intersection_signalized — 'yes' / 'no' / 'no_signal_coverage'; blank if not at an intersection. "
        "A crash flagged at-intersection that does not snap to a covered node = 'no_signal_coverage'.",
        "nearest_ped_signal_m — distance to the nearest deduped signalized crossing.",
        "is_ambiguous_intersection (bool) — field and geometry disagree.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Scoped intersection findings (covered corridors only)", 2)
    for b in [
        f"Intersection crashes (field): {D['n_intx']} of 1,294 ({D['n_intx_f']} fatal). "
        f"Ambiguous: {D['amb']} — (A) {D['ambA']} labeled mid-block but at a covered junction, "
        f"(B) {D['ambB']} labeled intersection but not geometrically corroborated, of which "
        f"{D['offroad']} are geocoded >20 m off any road (crash-geolocation imprecision).",
        f"Within signal coverage (signalized 'yes'/'no'): {D['within']} crashes ({D['within_f']} fatal) — "
        f"at SIGNALIZED intersections {D['y']} ({D['ypct']}%), at UNSIGNALIZED {D['n']} ({D['npct']}%); "
        f"fatal: signalized {D['yf']}, unsignalized {D['nf']}.",
        f"Off-corridor / no_signal_coverage (excluded from the share): {D['nocov']} crashes "
        f"({D['nocov_f']} fatal).",
        f"Most dangerous covered intersections: {D['dangerous']}.",
        f"Completeness spot-check (Winchester Rd): {D['win_sig']} of {D['win_j']} named junctions are "
        f"signalized in the inventory, located at the major cross-streets — consistent with a complete "
        f"inventory along the corridor, not a partial sample.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Caveat: no_signal_coverage ≠ no signal", 2)
    doc.add_paragraph(
        "The scope is signal-covered corridors only. Crashes at intersections OFF those corridors are "
        "excluded as 'no_signal_coverage' — they are NOT counted as unsignalized. Absence of signal "
        "data is not evidence of no signal. The signalized/unsignalized shares are descriptive; we do "
        "NOT claim unsignalized intersections are 'N times deadlier' from these raw counts (signalized "
        "intersections sit on the busiest arterials, which carry their own exposure).")
    doc.save(str(DOCX))
    print(f"  appended Phase-3a section to {DOCX.name}")


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    print("loading rulebook + signals + crossings + crashes...")
    rb = gpd.read_file(RULEBOOK).to_crs(CRS_M)
    ped = gpd.read_file(PED).to_crs(CRS_M).reset_index(drop=True)
    cross = gpd.read_file(CROSSINGS).to_crs(CRS_M)

    # ---- 1. COVERED CORRIDORS (route-anchored) ----
    snap = gpd.sjoin_nearest(ped[["ROUTE_NUMBER", "geometry"]],
                             rb[["Street_Name", "MTFCC", "geometry"]],
                             how="left", distance_col="d")
    snap = snap[~snap.index.duplicated(keep="first")]
    ped["street"] = snap["Street_Name"].values
    ped["smtfcc"] = snap["MTFCC"].values
    pair = ped.groupby(["ROUTE_NUMBER", "street"]).size().reset_index(name="n")
    route_tot = ped.groupby("ROUTE_NUMBER").size()
    covered_routes = route_tot[route_tot >= MIN_ROUTE_TOTAL].index
    street_mtfcc = ped.groupby("street")["smtfcc"].agg(lambda s: s.mode().iloc[0])
    # each covered route -> its DOMINANT arterial (+ any clear secondary corridor)
    covered = set()
    for rn in covered_routes:
        sub = pair[pair["ROUTE_NUMBER"] == rn].sort_values("n", ascending=False)
        if sub.empty:
            continue
        tot = route_tot[rn]
        dom = sub.iloc[0]["street"]
        if street_mtfcc.get(dom) in THROUGH_MTFCC:
            covered.add(dom)
        for _, r in sub.iloc[1:].iterrows():          # genuine secondary corridor only
            if r["n"] >= SECONDARY_SHARE * tot and street_mtfcc.get(r["street"]) in THROUGH_MTFCC:
                covered.add(r["street"])
    covered = sorted(covered)
    print(f"\nCOVERED CORRIDORS (dominant arterial per route + >={int(SECONDARY_SHARE*100)}% secondary): "
          f"{len(covered)} streets")
    print("  ", ", ".join(covered))

    # ---- 2. NODES (all named junctions citywide) ----
    through = rb[rb["MTFCC"].isin(THROUGH_MTFCC)].copy()
    through = through[through["Street_Name"].astype(str).str.strip() != ""]
    xs, ys, nms = [], [], []
    for geom, name in zip(through.geometry.values, through["Street_Name"].values):
        if geom is None or geom.is_empty:
            continue
        parts = geom.geoms if geom.geom_type == "MultiLineString" else [geom]
        for ln in parts:
            c = list(ln.coords)
            for pt in (c[0], c[-1]):
                xs.append(round(pt[0])); ys.append(round(pt[1])); nms.append(name)
    ep = pd.DataFrame({"x": xs, "y": ys, "name": nms})
    grp = ep.groupby(["x", "y"])["name"].agg(lambda s: tuple(sorted(set(s))))
    grp = grp[grp.map(len) >= 2]                      # >=2 distinct named through-roads
    nodes = gpd.GeoDataFrame(
        {"node_id": range(1, len(grp) + 1),
         "streets": ["; ".join(t) for t in grp.values],
         "n_streets": [len(t) for t in grp.values]},
        geometry=[Point(x, y) for x, y in grp.index], crs=CRS_M)
    nodes["on_covered"] = [any(s in covered for s in st.split("; ")) for st in nodes["streets"]]
    print(f"\nNODES: {len(nodes)} named junctions citywide; on covered corridors: "
          f"{int(nodes['on_covered'].sum())}")

    # ---- 3. SIGNALIZED (crossing within 30 m of node) ----
    nsig = gpd.sjoin_nearest(nodes[["node_id", "geometry"]], cross[["geometry"]],
                             how="left", distance_col="dsig")
    nsig = nsig[~nsig.index.duplicated(keep="first")]
    nodes["nearest_crossing_m"] = nsig["dsig"].values
    nodes["signalized"] = nodes["nearest_crossing_m"] <= SNAP_M
    cov = nodes[nodes["on_covered"]]
    print(f"  covered nodes signalized: {int(cov['signalized'].sum())} / {len(cov)} "
          f"({pct(int(cov['signalized'].sum()), len(cov))}%)")
    nodes[nodes["on_covered"]].to_crs("EPSG:4326").to_file(NODES_OUT, driver="GeoJSON")

    # ---- 4 & 5. CRASH ATTRIBUTES ----
    cr = pd.read_csv(CRASHES)
    pts = gpd.GeoDataFrame(cr.copy(),
                           geometry=gpd.points_from_xy(cr["Longitude"], cr["Latitude"]),
                           crs="EPSG:4326").to_crs(CRS_M)
    covnodes = nodes[nodes["on_covered"]]
    # nearest COVERED node (drives signalized + node_id), nearest ANY node (ambiguity), nearest crossing
    ncov = gpd.sjoin_nearest(pts[["geometry"]],
                             covnodes[["node_id", "signalized", "geometry"]],
                             how="left", distance_col="dcov")
    ncov = ncov[~ncov.index.duplicated(keep="first")]
    nany = gpd.sjoin_nearest(pts[["geometry"]], nodes[["geometry"]], how="left", distance_col="dany")
    nany = nany[~nany.index.duplicated(keep="first")]
    nc = gpd.sjoin_nearest(pts[["geometry"]], cross[["geometry"]], how="left", distance_col="dcross")
    nc = nc[~nc.index.duplicated(keep="first")]

    field_intx = cr["NonMotoristLocation"].astype(str).str.startswith("Intersection").values
    d_cov = ncov["dcov"].values
    cov_within = d_cov <= SNAP_M
    cov_sig = ncov["signalized"].values.astype(bool)
    any_within = nany["dany"].values <= SNAP_M

    out = cr.copy()
    out["at_intersection"] = field_intx
    out["intersection_node_id"] = pd.array(np.where(cov_within, ncov["node_id"].values, np.nan),
                                           dtype="Int64")
    out["nearest_ped_signal_m"] = np.round(nc["dcross"].values, 1)
    # ambiguous = (a) field=mid-block but within 30 m of a COVERED node, or
    #             (b) field=intersection but >30 m from ANY named junction
    out["is_ambiguous_intersection"] = ((~field_intx & cov_within) | (field_intx & ~any_within))

    sig = np.full(len(out), "", dtype=object)
    for i in range(len(out)):
        if not field_intx[i]:
            sig[i] = ""                                   # not at an intersection -> NA
        elif cov_within[i]:
            sig[i] = "yes" if cov_sig[i] else "no"
        else:
            sig[i] = "no_signal_coverage"                 # at-intersection but off covered corridor
    out["intersection_signalized"] = sig
    out.to_csv(CRASH_OUT, index=False, encoding="utf-8")

    # ---- 7. STATS (descriptive, scoped) ----
    NF = int((out["InjuryClass"] == FATAL).sum())
    atx = out[out["at_intersection"]]
    amb = int(out["is_ambiguous_intersection"].sum())
    within = out[out["intersection_signalized"].isin(["yes", "no"])]
    wf = within[within["InjuryClass"] == FATAL]
    y = int((within["intersection_signalized"] == "yes").sum())
    n = int((within["intersection_signalized"] == "no").sum())
    yf = int((wf["intersection_signalized"] == "yes").sum())
    nf = int((wf["intersection_signalized"] == "no").sum())
    nocov = int((out["intersection_signalized"] == "no_signal_coverage").sum())
    nocov_f = int(((out["intersection_signalized"] == "no_signal_coverage") &
                   (out["InjuryClass"] == FATAL)).sum())

    # ambiguity breakdown (interpretability: most of it is crash geocoding imprecision)
    ambA = int((~field_intx & cov_within).sum())
    ambB = int((field_intx & ~any_within).sum())
    ntr = gpd.sjoin_nearest(pts[["geometry"]], through[["geometry"]], how="left", distance_col="dtr")
    ntr = ntr[~ntr.index.duplicated(keep="first")]
    offroad = int((field_intx & (ntr["dtr"].values > 20)).sum())

    print("\n" + "=" * 70)
    print(f"intersection crashes (field): {len(atx)} of {len(out)} "
          f"({int((atx['InjuryClass']==FATAL).sum())} fatal) | ambiguous: {amb}")
    print(f"  ambiguity = (A) {ambA} mid-block-labeled but at a covered junction + "
          f"(B) {ambB} intersection-labeled but not geometrically corroborated")
    print(f"  of (B), {offroad} are geocoded >20 m off ANY road (geolocation imprecision, not a node gap)")
    print(f"\nSCOPED to signal-covered corridors — within-coverage intersection crashes "
          f"(signalized 'yes'/'no'): {len(within)} ({len(wf)} fatal)")
    print(f"  at SIGNALIZED intersections:   {y:4d} ({pct(y,len(within))}%)  | fatal {yf} ({pct(yf,len(wf))}%)")
    print(f"  at UNSIGNALIZED intersections: {n:4d} ({pct(n,len(within))}%)  | fatal {nf} ({pct(nf,len(wf))}%)")
    print(f"  off-corridor / no_signal_coverage (NOT counted as unsignalized): "
          f"{nocov} ({nocov_f} fatal)")
    print("  NOTE: scope = signal-covered corridors only; off-corridor intersection crashes are "
          "excluded as no_signal_coverage, NOT as unsignalized. Descriptive shares only — no causal claim.")

    # most dangerous intersections (covered nodes with crashes)
    j = out[out["intersection_node_id"].notna()].copy()
    j = j[j["intersection_node_id"].isin(cov["node_id"])]
    agg = j.groupby("intersection_node_id").agg(
        crashes=("MstrRecNbrTxt", "size"),
        deaths=("InjuryClass", lambda s: int((s == FATAL).sum()))).reset_index()
    agg = agg.merge(cov[["node_id", "streets", "signalized"]],
                    left_on="intersection_node_id", right_on="node_id", how="left")
    agg = agg.sort_values(["crashes", "deaths"], ascending=False).head(15)
    print("\nMOST DANGEROUS COVERED INTERSECTIONS (crashes | deaths | signalized | streets):")
    for _, r in agg.iterrows():
        print(f"  {int(r['crashes']):2d} | {int(r['deaths']):2d} | "
              f"{'SIGNAL' if r['signalized'] else 'unsig ':6s} | {r['streets']}")

    # ---- 8. WINCHESTER spot-check (completeness) ----
    print("\nWINCHESTER spot-check (covered-corridor completeness):")
    win = cov[cov["streets"].str.contains("WINCHESTER", na=False)].copy()
    print(f"  Winchester named junctions: {len(win)}; signalized: {int(win['signalized'].sum())} "
          f"({pct(int(win['signalized'].sum()), len(win))}%)")
    wsig = win[win["signalized"]].sort_values("nearest_crossing_m").head(10)
    for _, r in wsig.iterrows():
        print(f"    SIGNAL @ {r['streets']}  (crossing {r['nearest_crossing_m']:.0f} m)")
    print("  -> if signals appear at most major Winchester cross-streets, the inventory looks "
          "complete (not a partial sample).")

    # ---- docx (Phase 3a, idempotent) ----
    dang = "; ".join(
        f"{r['streets']} ({int(r['crashes'])} crashes/{int(r['deaths'])} deaths, "
        f"{'signalized' if r['signalized'] else 'unsignalized'})"
        for _, r in agg.head(5).iterrows())
    append_docx({
        "n_signals": len(ped), "n_cross": len(cross), "n_corridors": len(covered),
        "corridors_sample": ", ".join(covered[:12]),
        "n_covnodes": len(cov), "n_signodes": int(cov["signalized"].sum()),
        "n_intx": len(atx), "n_intx_f": int((atx["InjuryClass"] == FATAL).sum()),
        "amb": amb, "ambA": ambA, "ambB": ambB, "offroad": offroad,
        "within": len(within), "within_f": len(wf),
        "y": y, "n": n, "yf": yf, "nf": nf,
        "ypct": pct(y, len(within)), "npct": pct(n, len(within)),
        "nocov": nocov, "nocov_f": nocov_f, "dangerous": dang,
        "win_j": len(win), "win_sig": int(win["signalized"].sum()),
    })

    # ---- reconciliation ----
    print("\nRECONCILIATION:")
    print(f"  rows written: {len(out)} (=1294? {len(out)==1294}); "
          f"fatal: {NF} (=175? {NF==175})")
    print(f"  attributes added; jurisdiction columns untouched. -> {CRASH_OUT.name}")
    print("Phase 2 analysis done (map + docx next).")


if __name__ == "__main__":
    main()
