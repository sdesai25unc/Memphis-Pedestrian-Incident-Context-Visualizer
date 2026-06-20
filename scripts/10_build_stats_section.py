r"""
10_build_stats_section.py
========================

(a) Computes a new lighting statistic and (b) appends a STATS / FINDINGS
dashboard section to the BOTTOM of the existing interactive map page, turning
outputs/interactive_map/index.html into one map-plus-dashboard resource.

It does NOT rebuild the map (script 09 does that). It reads the current
index.html, leaves the map fully intact, and appends the dashboard below it.
Re-running is safe: it strips any previously-appended stats section first, so it
never double-appends. (Run order: 09 builds the map, then 10 adds the dashboard.)

All displayed numbers are COMPUTED from the data files - nothing is hardcoded.
Charts use Chart.js from a CDN (same approach as the map's Leaflet CDN), with the
chart data embedded as JS variables so the page still opens on file://.

Run it with:
    .\.venv\Scripts\python.exe scripts\10_build_stats_section.py
"""

import json
import datetime
from pathlib import Path

import pandas as pd
from docx import Document


# ---------------------------------------------------------------------------
# SETTINGS
# ---------------------------------------------------------------------------
FATAL_VALUE = "Fatal"
SERIOUS_VALUE = "Suspected Serious Injury"
COLOR_CITY = "#0e8f8f"   # teal
COLOR_TDOT = "#d62728"   # red

CHARTJS_CDN = "https://cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js"

PROJECT_ROOT = Path(__file__).resolve().parent.parent
PROCESSED = PROJECT_ROOT / "data" / "processed"
NAMED_CSV = PROCESSED / "shelby_crashes_named.csv"
DEADLIEST_CSV = PROCESSED / "deadliest_streets.csv"
DOCX_PATH = PROCESSED / "novel_statistics.docx"
HTML_PATH = PROJECT_ROOT / "outputs" / "interactive_map" / "index.html"

# Known street-type tokens, to flag a "name" that is really just a type.
TYPE_TOKENS = {
    "ALLEY", "ALY", "AVE", "BLVD", "ST", "RD", "DR", "LN", "CT", "CIR", "PL",
    "PKWY", "HWY", "WAY", "TER", "TRL", "PASS", "LOOP", "PIKE", "COVE", "CV",
}

# Run-once markers so re-runs replace (not duplicate) the appended section.
HEAD_MARKER = "/* stats-layout-override v1 */"
START_MARKER = "<!-- STATS_SECTION_START -->"
END_MARKER = "<!-- STATS_SECTION_END -->"


def pct(part, whole):
    return round(100.0 * part / whole, 1) if whole else 0.0


# ===========================================================================
# COMPUTE
# ===========================================================================
def compute():
    n = pd.read_csv(NAMED_CSV)
    s = pd.read_csv(DEADLIEST_CSV)
    fatal = n[n["InjuryClass"] == FATAL_VALUE]
    N, NF = len(n), len(fatal)

    out = {"n_all": N, "n_fatal": NF}

    # --- LIGHTING (Step 1) -------------------------------------------------
    def is_dark(series):
        return series.astype(str).str.startswith("Dark")

    def is_unlit(series):
        return series.astype(str).eq("Dark-Not Lighted")

    lighting = {}
    for label, df in [("all", n), ("fatal", fatal)]:
        tot = len(df)
        dark = int(is_dark(df["LightCondition"]).sum())
        unlit = int(is_unlit(df["LightCondition"]).sum())
        lighting[label] = {
            "n": tot,
            "dark": dark, "dark_pct": pct(dark, tot),
            "unlit": unlit, "unlit_pct": pct(unlit, tot),
            "breakdown": df["LightCondition"].value_counts(dropna=False).to_dict(),
        }
    out["lighting"] = lighting

    # --- JURISDICTION SPLIT ------------------------------------------------
    city_all = int((n["Jurisdiction"] == "City of Memphis").sum())
    tdot_all = int((n["Jurisdiction"] == "TDOT").sum())
    city_fatal = int((fatal["Jurisdiction"] == "City of Memphis").sum())
    tdot_fatal = int((fatal["Jurisdiction"] == "TDOT").sum())
    out["juris"] = {
        "city_all": city_all, "tdot_all": tdot_all,
        "city_all_pct": pct(city_all, N), "tdot_all_pct": pct(tdot_all, N),
        "city_fatal": city_fatal, "tdot_fatal": tdot_fatal,
        "city_fatal_pct": pct(city_fatal, NF), "tdot_fatal_pct": pct(tdot_fatal, NF),
    }

    # --- ROAD CHARACTER ----------------------------------------------------
    big40 = fatal["Street_SPDLIMIT"] >= 40
    lane4 = fatal["Street_LANES"] >= 4
    both = int((big40 & lane4).sum())
    out["design"] = {
        "fatal_both_pct": pct(both, NF), "fatal_both_n": both,
        "fatal_40_pct": pct(int(big40.sum()), NF),
        "fatal_4ln_pct": pct(int(lane4.sum()), NF),
    }
    # fatal crashes by lane count (1..max) for the chart
    lane_counts = fatal["Street_LANES"].value_counts().sort_index()
    out["lanes_chart"] = {
        "labels": [int(k) for k in lane_counts.index],
        "values": [int(v) for v in lane_counts.values],
    }

    # --- CONCENTRATION -----------------------------------------------------
    by_fatal = s.sort_values(["Fatal_Crashes", "Total_Crashes"], ascending=False)
    cum = by_fatal["Fatal_Crashes"].cumsum()
    streets_half = int((cum < NF / 2).sum() + 1)
    zero_fatal = int((s["Fatal_Crashes"] == 0).sum())
    out["concentration"] = {
        "streets_half": streets_half,
        "n_streets": len(s),
        "zero_fatal": zero_fatal,
        "zero_fatal_pct": pct(zero_fatal, len(s)),
    }

    # --- YEAR TREND --------------------------------------------------------
    by_year = n.groupby("YearNmb")
    years = sorted(int(y) for y in n["YearNmb"].unique())
    out["year_chart"] = {
        "labels": years,
        "all": [int((n["YearNmb"] == y).sum()) for y in years],
        "fatal": [int(((n["YearNmb"] == y) & (n["InjuryClass"] == FATAL_VALUE)).sum())
                  for y in years],
    }

    # --- TOP 25 TABLE ------------------------------------------------------
    top = s.sort_values(["Total_Crashes", "Fatal_Crashes"], ascending=False).head(25)
    rows = []
    for _, r in top.iterrows():
        rows.append({
            "name": r["Street_Name"],
            "owner": r["Dominant_Jurisdiction"],
            "mixed": bool(r["Mixed_Jurisdiction"]),
            "total": int(r["Total_Crashes"]),
            "fatal": int(r["Fatal_Crashes"]),
            "serious": int(r["Serious_Injuries"]),
            "lanes": (None if pd.isna(r["LANES"]) else int(r["LANES"])),
            "spd": (None if pd.isna(r["SPDLIMIT"]) else int(r["SPDLIMIT"])),
        })
    out["top25"] = rows

    # --- BROKEN-NAME FLAGS (report only; do NOT fix) -----------------------
    flags = []
    names = [r["name"] for r in rows]
    for nm in names:
        toks = nm.split()
        if len(toks) == 1 and toks[0] in TYPE_TOKENS:
            flags.append((nm, "name is just a street-type token"))
        elif nm.startswith("INTERSTATE") or nm.startswith("I-"):
            flags.append((nm, "interstate/highway - categorically different from arterials"))
    # directional-split detection: same base name appears with >1 direction in top25
    def base(nm):
        t = nm.split()
        if t and t[0] in {"N", "S", "E", "W"}:
            t = t[1:]
        if t and t[-1] in {"N", "S", "E", "W"}:
            t = t[:-1]
        return " ".join(t)
    from collections import defaultdict
    groups = defaultdict(list)
    for nm in names:
        groups[base(nm)].append(nm)
    for b, members in groups.items():
        if len(members) > 1:
            flags.append((" / ".join(members), "possible directional split of one corridor"))
    out["flags"] = flags

    return out


# ===========================================================================
# HTML
# ===========================================================================
HEAD_OVERRIDE = """
""" + HEAD_MARKER + """
html, body { height: auto; min-height: 100%; }
#header { position: sticky; }
#map { position: relative !important; top: auto; bottom: auto; left: auto; right: auto; height: 86vh; }
"""

STATS_TEMPLATE = START_MARKER + r"""
<section id="stats">
  <style>
    #stats { background: #f4f6f7; color: #1a1a1a; padding: 28px 20px 48px;
             font-family: "Segoe UI", Roboto, Helvetica, Arial, sans-serif; }
    #stats .inner { max-width: 1100px; margin: 0 auto; }
    #stats h2 { font-size: 24px; color: #14303f; margin: 4px 0 4px; }
    #stats .sub { color: #4a5b63; margin: 0 0 22px; font-size: 14px; }
    #stats h3 { font-size: 18px; color: #14303f; margin: 34px 0 12px;
                border-bottom: 2px solid #d8dee1; padding-bottom: 6px; }
    .hero { display: grid; grid-template-columns: repeat(auto-fit, minmax(190px, 1fr)); gap: 14px; }
    .card { background: #fff; border-radius: 8px; padding: 16px; box-shadow: 0 1px 4px rgba(0,0,0,.12);
            border-top: 4px solid #14303f; }
    .card.city { border-top-color: __CITY__; }
    .card.tdot { border-top-color: __TDOT__; }
    .card.dark { border-top-color: #2b2b50; }
    .card.ext  { border-top-color: #b8862b; }
    .card .big { font-size: 30px; font-weight: 700; color: #14303f; line-height: 1.1; }
    .card .lab { font-size: 13px; color: #33444c; margin-top: 6px; }
    .card .tag { display: inline-block; font-size: 10px; font-weight: 700; letter-spacing: .04em;
                 text-transform: uppercase; color: #8a6d24; background: #faf0d7; padding: 2px 6px;
                 border-radius: 3px; margin-bottom: 6px; }
    .reframe { background: #fff; border-left: 4px solid __CITY__; padding: 14px 16px; border-radius: 4px;
               font-size: 15px; line-height: 1.6; margin: 20px 0 0; box-shadow: 0 1px 4px rgba(0,0,0,.08); }
    .charts { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 18px; }
    .chart-box { background: #fff; border-radius: 8px; padding: 14px; box-shadow: 0 1px 4px rgba(0,0,0,.12); }
    .chart-box h4 { margin: 0 0 8px; font-size: 14px; color: #14303f; }
    .chart-box .cap { font-size: 11.5px; color: #6a7980; margin: 8px 0 0; }
    table.deadliest { width: 100%; border-collapse: collapse; background: #fff; font-size: 13px;
                      box-shadow: 0 1px 4px rgba(0,0,0,.12); border-radius: 8px; overflow: hidden; }
    table.deadliest th { background: #14303f; color: #fff; text-align: left; padding: 9px 10px;
                         cursor: pointer; user-select: none; white-space: nowrap; }
    table.deadliest th:hover { background: #1d4257; }
    table.deadliest td { padding: 8px 10px; border-bottom: 1px solid #eaeef0; }
    table.deadliest tr:nth-child(even) td { background: #f7f9fa; }
    .pill { display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 11px; color: #fff; }
    .pill.city { background: __CITY__; } .pill.tdot { background: __TDOT__; }
    .varies { font-size: 10px; color: #8a6d24; display: block; }
    .foot { margin-top: 30px; font-size: 12.5px; color: #4a5b63; line-height: 1.6; }
    .foot b { color: #14303f; }
  </style>
  <div class="inner">
    <h2>Findings &mdash; who owns the deadly roads, and why people die on them</h2>
    <p class="sub">Computed from __NALL__ pedestrian &amp; non-motorist crashes inside the City of Memphis (__NFATAL__ fatal), __DMIN__ to __DMAX__.</p>

    <div class="hero">
      <div class="card city">
        <div class="big">__C1_BIG__</div>
        <div class="lab">of crashes are on <b>City of Memphis</b>&ndash;owned roads (the rest are TDOT state routes)</div>
      </div>
      <div class="card tdot">
        <div class="big">__C2_BIG__</div>
        <div class="lab">of pedestrian <b>deaths</b> are on roads that are <b>both 4+ lanes and 40+ mph</b> &mdash; the design-problem signature</div>
      </div>
      <div class="card">
        <div class="big">__C3_BIG__ streets</div>
        <div class="lab">account for <b>half of all pedestrian deaths</b>; __C3_ZERO__ of streets with any crash had <b>zero</b> deaths</div>
      </div>
      <div class="card dark">
        <div class="big">__C4_BIG__</div>
        <div class="lab">of deaths are on <b>dark, unlit</b> roads; <b>__C4_DARK__</b> of all deaths happen after dark</div>
      </div>
      <div class="card ext">
        <span class="tag">External context</span>
        <div class="big">#1</div>
        <div class="lab">Memphis ranks <b>#1 nationally</b> in pedestrian fatality rate <span style="color:#8a6d24">(Smart Growth America, <i>Dangerous by Design 2024</i> &mdash; not our computed finding)</span></div>
      </div>
    </div>

    <p class="reframe">Most Memphis pedestrian deaths are not random residential accidents. They cluster on <b>wide, fast, multi-lane arterials</b> &mdash; __R_4LN__ of deaths are on roads with four or more lanes and __R_40__ on roads posted 40 mph or higher. When a road is built to move cars quickly through many lanes, a person trying to cross has little chance. This is a <b>design problem, not a behavior problem</b>: the conditions that kill pedestrians are set by how the roads are built and lit, which the City and TDOT control.</p>

    <h3>Charts</h3>
    <div class="charts">
      <div class="chart-box"><h4>Who owns the road &mdash; all crashes vs. deaths</h4><canvas id="chartJuris" height="220"></canvas><p class="cap">City of Memphis (teal) vs. TDOT state routes (red).</p></div>
      <div class="chart-box"><h4>Pedestrian deaths by number of lanes</h4><canvas id="chartLanes" height="220"></canvas><p class="cap">Each fatal crash counted on its nearest street's lane count &mdash; deaths rise on bigger roads.</p></div>
      <div class="chart-box"><h4>Crashes by year</h4><canvas id="chartYear" height="220"></canvas><p class="cap"><b>Read with caution:</b> the data window is short (~3 years) and 2026 is year-to-date; road conditions have not changed, so do not read a trend as success.</p></div>
    </div>

    <h3>Top 25 deadliest streets</h3>
    <table class="deadliest" id="deadliest">
      <thead><tr>
        <th data-k="name" data-t="s">Street</th>
        <th data-k="owner" data-t="s">Owner</th>
        <th data-k="total" data-t="n">Total</th>
        <th data-k="fatal" data-t="n">Fatal</th>
        <th data-k="serious" data-t="n">Serious</th>
        <th data-k="lanes" data-t="n">Lanes</th>
        <th data-k="spd" data-t="n">Speed</th>
      </tr></thead>
      <tbody></tbody>
    </table>

    <div class="foot">
      <p><b>Method.</b> Crashes are TDOT SAFETY non-motorist records for Shelby County, 2023-01-01 to __DMAX__, deduplicated to one row per crash. Each crash was assigned to the City of Memphis or TDOT by a nearest-road spatial join (state-route layer, 30 m threshold) and to its nearest named street, all in EPSG:32136 (Tennessee, meters). All figures here are <b>in-Memphis crashes only</b> (Suburban-Shelby and bad-coordinate rows excluded) and are recomputed from the data, not hand-entered.</p>
      <p><b>Sources.</b> Crash data: Tennessee SAFETY MapServer (TDOT). Roads, state routes, city boundary, and street centerline: City of Memphis Public Works GIS. National ranking: Smart Growth America, <i>Dangerous by Design 2024</i>.</p>
    </div>
  </div>
</section>
<script src="__CHARTJS__"></script>
<script>
var STATS_JURIS = __JS_JURIS__;
var STATS_LANES = __JS_LANES__;
var STATS_YEAR  = __JS_YEAR__;
var STATS_TOP25 = __JS_TOP25__;
var SC_CITY = "__CITY__", SC_TDOT = "__TDOT__";

/* ---- charts ---- */
new Chart(document.getElementById("chartJuris"), {
  type: "bar",
  data: { labels: ["All crashes", "Fatal crashes"],
    datasets: [
      { label: "City of Memphis", backgroundColor: SC_CITY,
        data: [STATS_JURIS.city_all, STATS_JURIS.city_fatal] },
      { label: "TDOT state route", backgroundColor: SC_TDOT,
        data: [STATS_JURIS.tdot_all, STATS_JURIS.tdot_fatal] } ] },
  options: { responsive: true, scales: { x: { stacked: true }, y: { stacked: true, beginAtZero: true } },
             plugins: { legend: { position: "bottom" } } }
});
new Chart(document.getElementById("chartLanes"), {
  type: "bar",
  data: { labels: STATS_LANES.labels.map(function(l){return l+" lane"+(l==1?"":"s");}),
    datasets: [{ label: "Fatal crashes", backgroundColor: "#2b2b50", data: STATS_LANES.values }] },
  options: { responsive: true, plugins: { legend: { display: false } },
             scales: { y: { beginAtZero: true, title: { display: true, text: "fatal crashes" } } } }
});
new Chart(document.getElementById("chartYear"), {
  type: "bar",
  data: { labels: STATS_YEAR.labels,
    datasets: [
      { label: "All crashes", backgroundColor: "#9bbcc7", data: STATS_YEAR.all },
      { label: "Fatal crashes", backgroundColor: SC_TDOT, data: STATS_YEAR.fatal } ] },
  options: { responsive: true, plugins: { legend: { position: "bottom" } },
             scales: { y: { beginAtZero: true } } }
});

/* ---- top-25 table (click-to-sort) ---- */
(function(){
  var tbody = document.querySelector("#deadliest tbody");
  function ownerCell(r){
    var cls = (r.owner === "TDOT") ? "tdot" : "city";
    var lab = (r.owner === "TDOT") ? "TDOT" : "City";
    var v = r.mixed ? "<span class='varies'>ownership varies</span>" : "";
    return "<span class='pill "+cls+"'>"+lab+"</span>"+v;
  }
  function render(rows){
    tbody.innerHTML = rows.map(function(r){
      return "<tr><td>"+r.name+"</td><td>"+ownerCell(r)+"</td><td>"+r.total+"</td><td>"+
        r.fatal+"</td><td>"+r.serious+"</td><td>"+(r.lanes==null?"&mdash;":r.lanes)+"</td><td>"+
        (r.spd==null?"&mdash;":r.spd+" mph")+"</td></tr>";
    }).join("");
  }
  var rows = STATS_TOP25.slice(), sortKey = "total", asc = false;
  render(rows);
  document.querySelectorAll("#deadliest th").forEach(function(th){
    th.addEventListener("click", function(){
      var k = th.getAttribute("data-k"), t = th.getAttribute("data-t");
      if (k === sortKey) asc = !asc; else { sortKey = k; asc = (t === "s"); }
      rows.sort(function(a,b){
        var x=a[k], y=b[k];
        if (t === "n") { x = (x==null?-1:x); y = (y==null?-1:y); return asc ? x-y : y-x; }
        x=String(x); y=String(y); return asc ? x.localeCompare(y) : y.localeCompare(x);
      });
      render(rows);
    });
  });
})();
</script>
""" + END_MARKER + "\n"


def build_stats_html(d):
    j = d["juris"]
    html = STATS_TEMPLATE
    repl = {
        "__CITY__": COLOR_CITY, "__TDOT__": COLOR_TDOT, "__CHARTJS__": CHARTJS_CDN,
        "__NALL__": f"{d['n_all']:,}", "__NFATAL__": str(d["n_fatal"]),
        "__DMIN__": d["date_min"], "__DMAX__": d["date_max"],
        "__C1_BIG__": f"{j['city_all_pct']}%",
        "__C2_BIG__": f"{d['design']['fatal_both_pct']}%",
        "__C3_BIG__": str(d["concentration"]["streets_half"]),
        "__C3_ZERO__": f"{d['concentration']['zero_fatal_pct']}%",
        "__C4_BIG__": f"{d['lighting']['fatal']['unlit_pct']}%",
        "__C4_DARK__": f"{d['lighting']['fatal']['dark_pct']}%",
        "__R_4LN__": f"{d['design']['fatal_4ln_pct']}%",
        "__R_40__": f"{d['design']['fatal_40_pct']}%",
        "__JS_JURIS__": json.dumps(j, separators=(",", ":")),
        "__JS_LANES__": json.dumps(d["lanes_chart"], separators=(",", ":")),
        "__JS_YEAR__": json.dumps(d["year_chart"], separators=(",", ":")),
        "__JS_TOP25__": json.dumps(d["top25"], separators=(",", ":")).replace("<", "\\u003c").replace(">", "\\u003e"),
    }
    for k, v in repl.items():
        html = html.replace(k, v)
    return html


def update_html(d):
    html = HTML_PATH.read_text(encoding="utf-8")

    # 1. Make the layout scrollable (map becomes a sized block) - once.
    if HEAD_MARKER not in html:
        html = html.replace("</style>", HEAD_OVERRIDE + "</style>", 1)

    # 2. Strip any previously-appended stats section (idempotent re-run).
    if START_MARKER in html and END_MARKER in html:
        pre = html.split(START_MARKER)[0]
        post = html.split(END_MARKER)[1]
        html = pre + post

    # 3. Insert the fresh stats section just before </body>.
    section = build_stats_html(d)
    html = html.replace("</body>", section + "</body>", 1)
    HTML_PATH.write_text(html, encoding="utf-8")


def append_to_docx(d):
    if not DOCX_PATH.exists():
        print(f"  NOTE: {DOCX_PATH.name} not found - skipping docx append.")
        return
    doc = Document(str(DOCX_PATH))
    today = datetime.date.today().isoformat()
    lf, la = d["lighting"]["fatal"], d["lighting"]["all"]
    doc.add_heading(f"Lighting condition of crashes (added {today})", level=2)
    doc.add_paragraph(
        f"Of {la['n']:,} in-Memphis crashes, {la['dark']} ({la['dark_pct']}%) "
        f"occurred in a dark condition (Dark-Lighted, Dark-Not Lighted, or "
        f"Dark-Unknown Lighting); {la['unlit']} ({la['unlit_pct']}%) on a "
        f"specifically dark, UNLIT road."
    )
    doc.add_paragraph(
        f"Of {lf['n']} fatal in-Memphis crashes, {lf['dark']} ({lf['dark_pct']}%) "
        f"happened after dark and {lf['unlit']} ({lf['unlit_pct']}%) on a dark, "
        f"UNLIT road - the design-failure case (a road the City/State never lit). "
        f"Dawn and dusk are excluded from 'dark.'"
    )
    doc.add_paragraph(
        f"[{today}] Built the interactive stats/findings dashboard onto the bottom "
        f"of outputs/interactive_map/index.html (script 10)."
    )
    doc.save(str(DOCX_PATH))


def main():
    print("Computing stats + building dashboard section...")
    d = compute()
    # date range for captions
    n = pd.read_csv(NAMED_CSV)
    d["date_min"], d["date_max"] = n["CollisionDate"].min(), n["CollisionDate"].max()

    lf, la = d["lighting"]["fatal"], d["lighting"]["all"]
    print("\n" + "=" * 64)
    print("LIGHTING (computed)")
    print("=" * 64)
    print("Distinct LightCondition values (all in-Memphis):")
    for k, v in la["breakdown"].items():
        print(f"    {k}: {v}")
    print(f"\n  ALL  (n={la['n']}):   any-Dark {la['dark']} ({la['dark_pct']}%) | "
          f"Dark-Not Lighted {la['unlit']} ({la['unlit_pct']}%)")
    print(f"  FATAL (n={lf['n']}):   any-Dark {lf['dark']} ({lf['dark_pct']}%) | "
          f"Dark-Not Lighted {lf['unlit']} ({lf['unlit_pct']}%)")

    j, des, con = d["juris"], d["design"], d["concentration"]
    print("\n" + "=" * 64)
    print("HERO CARD NUMBERS (computed)")
    print("=" * 64)
    print(f"  1. City-owned share (all crashes):      {j['city_all_pct']}%  ({j['city_all']}/{d['n_all']})")
    print(f"  2. Fatal on BOTH 4+ lanes & 40+ mph:    {des['fatal_both_pct']}%  ({des['fatal_both_n']}/{d['n_fatal']})")
    print(f"  3. Streets = half of deaths:            {con['streets_half']}    | zero-fatal streets {con['zero_fatal_pct']}% ({con['zero_fatal']}/{con['n_streets']})")
    print(f"  4. Fatal on dark UNLIT roads:           {lf['unlit_pct']}%  | any-dark {lf['dark_pct']}%")
    print(f"  5. #1 nationally (EXTERNAL - SGA 2024)")
    print(f"  (reframe: fatal >=4 lanes {des['fatal_4ln_pct']}% | fatal >=40 mph {des['fatal_40_pct']}%)")

    print("\n" + "=" * 64)
    print("BROKEN-NAME FLAGS in top-25 (review for a later cleanup; NOT fixed here)")
    print("=" * 64)
    if d["flags"]:
        for nm, why in d["flags"]:
            print(f"  - {nm}  ->  {why}")
    else:
        print("  (none detected)")

    update_html(d)
    append_to_docx(d)

    print("\n" + "=" * 64)
    print(f"Updated map+dashboard -> {HTML_PATH}")
    print(f"Appended lighting stat + note -> {DOCX_PATH.name}")
    print("NOTE: open index.html by double-click; needs internet for the Leaflet +")
    print("      Chart.js CDNs and OpenStreetMap tiles.")
    print("=" * 64)


if __name__ == "__main__":
    main()
