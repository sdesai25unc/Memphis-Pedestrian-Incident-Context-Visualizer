r"""
23_union_poc.py
==============

Phase 3b — UNION AVE distance-to-crossing PROOF OF CONCEPT (Union only).

Builds a combined safe-crossing inventory on the in-Memphis stretch of Union Ave
(OSM marked crosswalks + TDOT pedestrian signals), measures each pedestrian crash's
ALONG-CORRIDOR distance to the nearest safe crossing, reports the corridor's longest
gap vs the FHWA ~300 ft best-practice spacing, and writes a focused Union map.

Signed-off method choices:
  - distance = along-corridor (linear referencing on the Union centerline);
  - stat computed on the 36 CROSSING-RELEVANT crashes (NonMotorist location = On Roadway
    / In Crosswalk); the 5 non-crossing (Outside Traffic / Not on Roadway / Unknown) reported separately;
  - SIGNALIZED safe crossing = TDOT signal only (OSM supplies MARKED crosswalks);
  - SAFE CROSSING = OSM marked crosswalk OR TDOT signal within 30 m of the centerline,
    deduped so a signalized intersection that also has a marked crosswalk = ONE crossing.

UNION ONLY — does not touch the citywide map or add OSM crossings elsewhere.
New files only.

Run it with:
    .\.venv\Scripts\python.exe scripts\23_union_poc.py
"""

import sys
import json
import math
import time
from pathlib import Path

import requests
import numpy as np
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point, LineString
from shapely.ops import linemerge, unary_union, substring

ROOT = Path(__file__).resolve().parent.parent
RAW = ROOT / "data" / "raw"
PROC = ROOT / "data" / "processed"
OUTMAP = ROOT / "outputs" / "interactive_map"
OSM_UNION = RAW / "osm_union_crossings.geojson"
RULEBOOK = PROC / "road_ownership_rulebook.geojson"
BOUNDARY = RAW / "memphis_boundary.geojson"
TDOT_DEDUP = PROC / "signalized_crossings_dedup.geojson"
PED_RAW = RAW / "ped_signals.geojson"
CRASHES = PROC / "shelby_crashes_signals.csv"
VIZ = OUTMAP / "union_poc.html"
REPORT = ROOT / "outputs" / "union_poc_report.md"
DOCX = PROC / "novel_statistics.docx"

CRS_M, CRS_GEO = "EPSG:32136", "EPSG:4326"
NEAR_M = 30.0            # safe crossing / crash within this of the Union centerline
DEDUP_NW = 8.0          # OSM node+way duplicate merge
DEDUP_SAFE = 30.0       # one safe crossing per intersection
FT = 0.3048
T100, T250, FHWA = 100 * FT, 250 * FT, 300 * FT
MIRRORS = ["https://overpass.kumi.systems/api/interpreter",
           "https://maps.mail.ru/osm/tools/overpass/api/interpreter",
           "https://overpass-api.de/api/interpreter"]


def overpass(q):
    last = None
    for ep in MIRRORS:
        try:
            r = requests.post(ep, data={"data": q}, timeout=180,
                              headers={"User-Agent": "memphis-ped-union-poc/1.0"})
            if r.status_code == 200 and r.headers.get("content-type", "").startswith("application/json"):
                return r.json()
            last = f"{ep}->{r.status_code}"
        except Exception as e:
            last = f"{ep}->{repr(e)[:60]}"
        time.sleep(1)
    raise RuntimeError(f"Overpass failed: {last}")


def acquire_union_osm(ref_geo):
    """Pull OSM crossings near Union WITH full geometry (out geom -> lines for ways)."""
    if OSM_UNION.exists():
        print(f"  using cached {OSM_UNION.name} (delete to re-pull)")
        return gpd.read_file(OSM_UNION)
    minx, miny, maxx, maxy = ref_geo.bounds
    pad = 0.004  # ~400 m
    s, w, n, e = miny - pad, minx - pad, maxy + pad, maxx + pad
    q = f"""[out:json][timeout:120];
(
  node["highway"="crossing"]({s},{w},{n},{e});
  node["crossing"]({s},{w},{n},{e});
  way["footway"="crossing"]({s},{w},{n},{e});
);
out geom;"""
    js = overpass(q)
    rows, geoms = [], []
    for el in js.get("elements", []):
        t = el.get("tags", {})
        if el["type"] == "node":
            g = Point(el["lon"], el["lat"])
        else:
            coords = [(p["lon"], p["lat"]) for p in el.get("geometry", [])]
            if len(coords) < 2:
                continue
            g = LineString(coords)
        rows.append({"osm_type": el["type"], "osm_id": el["id"],
                     "highway": t.get("highway", ""), "crossing": t.get("crossing", ""),
                     "markings": t.get("crossing:markings", ""), "footway": t.get("footway", "")})
        geoms.append(g)
    gdf = gpd.GeoDataFrame(rows, geometry=geoms, crs=CRS_GEO)
    gdf.to_file(OSM_UNION, driver="GeoJSON")
    return gdf


def uf(idx_pairs, n):
    parent = list(range(n))
    def find(a):
        while parent[a] != a:
            parent[a] = parent[parent[a]]; a = parent[a]
        return a
    for a, b in idx_pairs:
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[ra] = rb
    return np.array([find(i) for i in range(n)])


def cluster_pts(gdf, radius):
    g = gdf.reset_index(drop=True)
    if not len(g):
        return np.array([], dtype=int)
    buf = gpd.GeoDataFrame(geometry=g.buffer(radius), crs=g.crs)
    buf["bi"] = range(len(buf))
    g2 = g.copy(); g2["pi"] = range(len(g))
    j = gpd.sjoin(g2[["pi", "geometry"]], buf[["bi", "geometry"]], predicate="intersects", how="inner")
    return uf(list(zip(j["pi"], j["bi"])), len(g))


def _bear(p, q):
    return math.degrees(math.atan2(q[1] - p[1], q[0] - p[0])) % 180


def _ref_bear(ref, pt):
    a = ref.project(pt)
    p0 = ref.interpolate(max(0, a - 8)); p1 = ref.interpolate(min(ref.length, a + 8))
    return _bear((p0.x, p0.y), (p1.x, p1.y))


def _angle_to_ref(line, ref):
    """Angle (deg) between a crosswalk line and Union at the nearest point. 0=parallel, 90=perpendicular."""
    cs = list(line.coords)
    cb = _bear(cs[0], cs[-1])
    mid = line.interpolate(0.5, normalized=True)
    a = abs(cb - _ref_bear(ref, mid))
    return min(a, 180 - a)


def fmt_ft(m):
    return f"{m/FT:.0f} ft ({m:.0f} m)"


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    L = []
    def log(s=""):
        print(s); L.append(s)

    # ---- STEP 1: Union centerline reference ----
    rb = gpd.read_file(RULEBOOK).to_crs(CRS_M)
    bnd = gpd.read_file(BOUNDARY).to_crs(CRS_M).union_all()
    u = rb[(rb["Street_Name"] == "UNION AVE")]
    u = u[u.intersects(bnd)]
    merged = linemerge(unary_union(u.geometry.values))
    comps = list(merged.geoms) if merged.geom_type == "MultiLineString" else [merged]
    ref = max(comps, key=lambda l: l.length)            # single reference line
    ref_mi = ref.length / 1609.344
    ref_geo = gpd.GeoSeries([ref], crs=CRS_M).to_crs(CRS_GEO).iloc[0]

    log("# Union Ave — distance-to-crossing proof of concept\n")
    log("*Union only. Distances along the centerline (EPSG:32136). Safe crossing = OSM marked "
        "crosswalk OR TDOT pedestrian signal within 30 m, deduped (co-located = one).*\n")
    log(f"## 1. Corridor\n- In-Memphis Union Ave reference centerline: **{ref_mi:.2f} mi** "
        f"({ref.length:.0f} m), single carriageway.")

    # ---- STEP 0/2: OSM marked + TDOT signals near Union ----
    osm = acquire_union_osm(ref_geo).to_crs(CRS_M)
    mk = osm["markings"].astype(str).str.lower()
    osm["is_marked"] = ~mk.isin(["", "no", "nan", "none"])
    marked = osm[osm["is_marked"]].copy()
    marked = marked[marked.distance(ref) <= NEAR_M].copy()

    # Union junctions (where Union crosses a named cross-street) + cross-street geometry,
    # used to keep only crossings that cross UNION itself (not a side street).
    cross = rb[(rb["MTFCC"].isin(["S1200", "S1400"])) & (rb["Street_Name"] != "UNION AVE")]
    cross = cross[cross.distance(ref) <= 200]
    cross_u = unary_union(cross.geometry.values)
    jx = ref.intersection(unary_union(cross[cross.distance(ref) <= 5].geometry.values))
    jpts = gpd.GeoSeries([g for g in (jx.geoms if jx.geom_type.startswith("Multi") else [jx])
                          if g.geom_type == "Point"], crs=CRS_M)

    # decide per deduped marked location whether it crosses Union
    marked["rep"] = marked.geometry.centroid
    mrep = gpd.GeoDataFrame(marked.drop(columns="geometry"), geometry=marked["rep"], crs=CRS_M)
    marked["mloc"] = cluster_pts(mrep, DEDUP_NW)
    keep_reps, viz_geoms = [], []
    drop_line = drop_pt = 0
    for _, grp in marked.groupby("mloc"):
        lines = [g for g in grp.geometry if g.geom_type == "LineString"]
        if lines:                                              # line -> perpendicular & crosses Union
            ln = max(lines, key=lambda l: l.length)
            ok = (_angle_to_ref(ln, ref) > 45) and (ln.intersects(ref) or ln.distance(ref) <= 12)
            rep, vg = ln.interpolate(0.5, normalized=True), ln
            drop_line += 0 if ok else 1
        else:                                                  # point-only -> nearest centerline is Union
            pt = grp.geometry.iloc[0]
            d_ref = pt.distance(ref)
            ok = (d_ref <= 8) and (d_ref <= (pt.distance(cross_u) if not cross_u.is_empty else 1e9))
            rep, vg = pt, pt
            drop_pt += 0 if ok else 1
        if ok:
            keep_reps.append(rep); viz_geoms.append(vg)
    marked_keep = gpd.GeoDataFrame(geometry=keep_reps, crs=CRS_M)
    marked_viz = gpd.GeoDataFrame(geometry=viz_geoms, crs=CRS_M)

    # TDOT signals: keep only those at a junction ON Union
    tdot_all = gpd.read_file(TDOT_DEDUP).to_crs(CRS_M)
    tdot_all = tdot_all[tdot_all.distance(ref) <= NEAR_M].copy()
    if len(jpts):
        tdot = tdot_all[tdot_all.geometry.apply(lambda p: jpts.distance(p).min() <= NEAR_M)].copy()
    else:
        tdot = tdot_all.copy()
    drop_sig = len(tdot_all) - len(tdot)
    log(f"\n*Union-crossing filter: dropped {drop_line} parallel/side-street marked LINES, "
        f"{drop_pt} off-Union marked points, and {drop_sig} side-street-only TDOT signals "
        f"(kept only crossings that cross Union itself).*")

    # ---- combine -> safe crossings (signalized vs marked-only) ----
    comb = gpd.GeoDataFrame(
        {"src": ["marked"] * len(marked_keep) + ["signal"] * len(tdot)},
        geometry=list(marked_keep.geometry) + list(tdot.geometry), crs=CRS_M).reset_index(drop=True)
    clab = cluster_pts(comb, DEDUP_SAFE)
    comb["c"] = clab
    safe = []
    for c, g in comb.groupby("c"):
        has_sig = (g["src"] == "signal").any()
        cen = unary_union(g.geometry.values).centroid
        safe.append({"type": "signalized" if has_sig else "marked-only", "geometry": cen})
    safe = gpd.GeoDataFrame(safe, crs=CRS_M)
    safe["along"] = safe.geometry.apply(lambda p: ref.project(p))
    safe = safe.sort_values("along").reset_index(drop=True)
    n_sig = int((safe["type"] == "signalized").sum())
    n_mk = int((safe["type"] == "marked-only").sum())

    log(f"\n## 2. Combined safe-crossing inventory ({len(safe)} crossings)\n")
    log(f"- **Signalized (TDOT): {n_sig}** | **Marked-only (OSM): {n_mk}**  "
        f"(after the Union-crossing filter: {len(marked_keep)} Union-crossing marked locations + "
        f"{len(tdot)} Union signals, deduped; raw marked features near Union were {len(marked)}).")
    # spacing between consecutive safe crossings
    al = safe["along"].values
    spac = np.diff(al)
    if len(spac):
        log(f"- Spacing between consecutive safe crossings: median **{fmt_ft(np.median(spac))}**, "
            f"mean {fmt_ft(spac.mean())}, max {fmt_ft(spac.max())}.")

    # ---- STEP 3: Union crashes ----
    cdf = pd.read_csv(CRASHES)
    cg = gpd.GeoDataFrame(cdf, geometry=gpd.points_from_xy(cdf.Longitude, cdf.Latitude),
                          crs=CRS_GEO).to_crs(CRS_M)
    cg["d_ref"] = cg.distance(ref)
    near = cg[cg["d_ref"] <= NEAR_M].copy()
    near["fatal"] = near["InjuryClass"] == "Fatal"
    loc = near["NonMotoristLocation"].astype(str)
    near["crossing_rel"] = loc.str.contains("On Roadway", case=True) | loc.str.contains("In Crosswalk", case=True)
    rel = near[near["crossing_rel"]].copy()
    nonrel = near[~near["crossing_rel"]]
    log(f"\n## 3. Union crashes (within {NEAR_M:.0f} m of centerline)\n")
    log(f"- Total: **{len(near)}** ({int(near['fatal'].sum())} fatal)  *(deadliest-list anchor ≈ 36/8; "
        f"the 30 m buffer catches a few more than the nearest-street assignment)*.")
    log(f"- **Crossing-relevant (On Roadway / In Crosswalk): {len(rel)}** "
        f"({int(rel['fatal'].sum())} fatal) — the headline set.")
    log(f"- Non-crossing (Outside Traffic / Not on Roadway / Unknown): {len(nonrel)} "
        f"({int(nonrel['fatal'].sum())} fatal) — reported, excluded from the distance stat.")

    # ---- STEP 4: along-corridor distance to nearest safe crossing ----
    rel["along"] = rel.geometry.apply(lambda p: ref.project(p))
    if len(safe):
        rel["d_cross"] = rel["along"].apply(lambda a: float(np.min(np.abs(al - a))))
    else:
        rel["d_cross"] = np.nan
    d = rel["d_cross"]
    log(f"\n## 4. Distance from a crossing-relevant crash to the nearest safe crossing "
        f"(along-corridor)\n")
    log(f"- mean **{fmt_ft(d.mean())}**, median **{fmt_ft(d.median())}**, max {fmt_ft(d.max())}.")
    log(f"- struck **> 100 ft** from the nearest safe crossing: **{int((d>T100).sum())}/{len(rel)} "
        f"({100*(d>T100).mean():.0f}%)**; **> 250 ft**: **{int((d>T250).sum())}/{len(rel)} "
        f"({100*(d>T250).mean():.0f}%)**.")
    n_at = int((d <= T100).sum()); n_gap = int((d > T250).sum()); n_mid = len(rel) - n_at - n_gap
    log(f"- **Bimodal split:** **{n_at} struck AT/near a Union crossing** (≤100 ft), "
        f"**{n_gap} struck in a gap** (>250 ft), {n_mid} in between — of {len(rel)} crossing-relevant.")
    log(f"- fatal crossing-relevant crashes' distances: "
        f"{', '.join(fmt_ft(x) for x in sorted(rel.loc[rel['fatal'],'d_cross']))}.")

    # corridor gap (longest stretch with no safe crossing, incl. ends)
    bounds = np.concatenate([[0.0], al, [ref.length]])
    gaps = np.diff(bounds)
    gi = int(gaps.argmax())
    gap_len = gaps[gi]; gstart, gend = bounds[gi], bounds[gi + 1]
    log(f"\n## 5. Longest gap vs FHWA best-practice spacing\n")
    log(f"- Longest stretch of Union with **no safe crossing: {fmt_ft(gap_len)}** "
        f"(from {gstart/1609.344:.2f} to {gend/1609.344:.2f} mi along the corridor"
        f"{' — at the corridor end' if gi==0 or gi==len(gaps)-1 else ''}).")
    log(f"- Median safe-crossing spacing **{fmt_ft(np.median(spac)) if len(spac) else 'n/a'}** vs the "
        f"FHWA marked-crossing best-practice guidance of ~300 ft (91 m). The longest gap is "
        f"**{gap_len/FHWA:.1f}×** the ~300 ft figure. *(FHWA ~300 ft is best-practice spacing guidance, "
        f"not a legal standard.)*")

    # ---- machine-readable Union summary (consumed by the search index, script 24) ----
    unodes = gpd.read_file(PROC / "intersection_nodes_covered.geojson").to_crs(CRS_M)
    unodes = unodes[unodes["streets"].astype(str).str.contains("UNION AVE")].copy()
    unodes["nsc"] = (unodes.geometry.apply(lambda p: float(np.min(np.abs(al - ref.project(p)))))
                     if len(safe) else np.nan)
    summary = {
        "corridor": "UNION AVE", "ref_mi": round(ref_mi, 2),
        "n_safe": int(len(safe)), "n_signalized": int(n_sig), "n_marked_only": int(n_mk),
        "longest_gap_ft": round(gap_len / FT),
        "median_spacing_ft": (round(float(np.median(spac)) / FT) if len(spac) else None),
        "n_crash": int(len(near)), "n_fatal": int(near["fatal"].sum()),
        "n_crossing_rel": int(len(rel)), "mean_ft": round(float(d.mean()) / FT),
        "median_ft": round(float(d.median()) / FT),
        "pct_over_250ft": round(100 * float((d > T250).mean())), "n_over_250ft": int((d > T250).sum()),
        "node_nearest_safe_m": {int(r.node_id): round(float(r.nsc), 1)
                                for _, r in unodes.iterrows() if pd.notna(r.nsc)},
    }
    (PROC / "union_safe_summary.json").write_text(json.dumps(summary, indent=1), encoding="utf-8")

    # ---- STEP 5: focused Union viz ----
    write_viz(ref, ref_geo, safe, marked_viz, tdot, rel, nonrel, gstart, gend, n_sig, n_mk)
    log(f"\n## 6. Visualization\n- Focused Union map written to `{VIZ.relative_to(ROOT)}` "
        f"(crosswalk lines by type, TDOT signal points per-corner, crashes shaded by distance, "
        f"longest gap highlighted). The citywide map is untouched.")

    log("\n## Method judgment calls\n"
        "- Distance is along-corridor (linear referencing on the single Union reference line), not "
        "straight-line — it reflects walking distance along the road.\n"
        "- 'In Crosswalk' crashes (struck AT a marked crossing, distance ≈ 0) are kept in the "
        "crossing-relevant set; they pull the mean down but honestly show people are struck even at "
        "crossings.\n"
        "- Signalized crossings are TDOT-only (more complete than OSM signals); marked crosswalks are "
        "OSM, which Phase 3a found dense and continuous on Union specifically.\n"
        "- OSM was re-pulled with full line geometry for Union only (the citywide file stored points).")

    REPORT.write_text("\n".join(L) + "\n", encoding="utf-8")
    append_docx(ref_mi, len(safe), n_sig, n_mk, len(near), int(near['fatal'].sum()), len(rel),
                d.mean(), d.median(), int((d > T250).sum()), gap_len)
    print(f"\nWrote {OSM_UNION.name}, {VIZ.name}, {REPORT.name}; appended docx.")
    print("UNION ONLY — citywide map untouched. STOP for review.")


def write_viz(ref, ref_geo, safe, marked_viz, tdot, rel, nonrel, gstart, gend, n_sig, n_mk):
    def gj(geom_series_gdf):
        return json.loads(geom_series_gdf.to_crs(CRS_GEO).to_json())
    # kept (Union-crossing) marked crosswalks, tagged signalized vs marked-only by proximity to a signal
    sig_pts = safe[safe["type"] == "signalized"]
    marked = marked_viz.copy()
    marked["type"] = marked.geometry.centroid.apply(
        lambda p: "signalized" if (len(sig_pts) and sig_pts.distance(p).min() <= DEDUP_SAFE) else "marked-only")
    osm_gj = gj(marked[["type", "geometry"]])
    # raw per-corner TDOT signal heads/buttons, only at the KEPT Union signals
    praw = gpd.read_file(PED_RAW).to_crs(CRS_M)
    praw = praw[praw.distance(ref) <= NEAR_M]
    if len(tdot):
        praw = praw[praw.geometry.apply(lambda p: tdot.distance(p).min() <= NEAR_M)]
    tdot_gj = gj(praw[["FEATURE_DESCRIPTION", "geometry"]])
    safe_gj = gj(safe[["type", "geometry"]])
    # crashes with distance
    rel2 = rel.copy(); rel2["nonrel"] = False
    nonrel2 = nonrel.copy(); nonrel2["d_cross"] = np.nan; nonrel2["nonrel"] = True
    allc = pd.concat([rel2, nonrel2])
    allc_g = gpd.GeoDataFrame(
        {"d_ft": (allc["d_cross"] / FT).round(0), "fatal": allc["InjuryClass"] == "Fatal",
         "nonrel": allc["nonrel"], "date": allc["CollisionDate"].astype(str),
         "loc": allc["NonMotoristLocation"].astype(str)},
        geometry=allc.geometry, crs=CRS_M)
    crash_gj = gj(allc_g)
    gap_geo = gpd.GeoSeries([substring(ref, gstart, gend)], crs=CRS_M).to_crs(CRS_GEO).iloc[0]
    ref_xy = [[y, x] for x, y in ref_geo.coords]
    c = ref_geo.centroid

    html = _TEMPLATE
    for k, v in {
        "__CENTER__": json.dumps([c.y, c.x]),
        "__REF__": json.dumps(ref_xy),
        "__GAP__": json.dumps([[p[1], p[0]] for p in gap_geo.coords]),
        "__OSM__": json.dumps(osm_gj), "__TDOT__": json.dumps(tdot_gj),
        "__SAFE__": json.dumps(safe_gj), "__CRASH__": json.dumps(crash_gj),
        "__NSIG__": str(n_sig), "__NMK__": str(n_mk),
    }.items():
        html = html.replace(k, v)
    OUTMAP.mkdir(parents=True, exist_ok=True)
    VIZ.write_text(html, encoding="utf-8")


def append_docx(ref_mi, n_safe, n_sig, n_mk, n_crash, n_fatal, n_rel, dmean, dmed, n250, gap):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print("  (python-docx not available; skipping docx append)")
        return
    doc = Document(DOCX) if DOCX.exists() else Document()
    # idempotent: drop any prior Union POC section (its heading to the end of the doc) before re-appending
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Union Ave —"):
            start = i; break
    if start is not None:
        for p in doc.paragraphs[start:]:
            p._element.getparent().remove(p._element)
    doc.add_heading("Union Ave — Distance-to-Crossing Proof of Concept (added 2026-06-17)", level=1)
    for line in [
        f"Scope: in-Memphis Union Ave reference centerline, {ref_mi:.2f} mi (single carriageway).",
        f"Safe-crossing inventory (OSM marked crosswalk OR TDOT pedestrian signal within 30 m, deduped "
        f"so a signalized intersection with a marked crosswalk = one): {n_safe} crossings = "
        f"{n_sig} signalized (TDOT) + {n_mk} marked-only (OSM).",
        f"Union crashes within 30 m: {n_crash} ({n_fatal} fatal); crossing-relevant (On Roadway / In "
        f"Crosswalk): {n_rel} — the set used for the distance stat.",
        f"Along-corridor distance from a crossing-relevant crash to the nearest safe crossing: mean "
        f"{dmean/FT:.0f} ft, median {dmed/FT:.0f} ft; {n250}/{n_rel} struck more than 250 ft from one.",
        f"Longest stretch of Union with no safe crossing: {gap/FT:.0f} ft ({gap/FHWA:.1f}x the FHWA "
        f"~300 ft marked-crossing best-practice spacing — best practice, not a legal standard).",
        "Caveats: PROOF OF CONCEPT on Union only (the one corridor Phase 3a found OSM-complete). "
        "Signalized = TDOT (more complete than OSM signals); marked = OSM. Distance is along-corridor "
        "linear referencing. 'In Crosswalk' crashes are kept (struck at a crossing, ~0 ft).",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(DOCX)
    print(f"  appended Union POC section to {DOCX.name}")


_TEMPLATE = r"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>Union Ave — pedestrian crashes vs safe crossings</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css">
<style>
  body{margin:0;font-family:system-ui,Segoe UI,Roboto,sans-serif}
  #map{height:100vh}
  .hdr{position:absolute;z-index:1000;top:10px;left:50px;background:#fff;padding:8px 12px;
       border-radius:6px;box-shadow:0 1px 6px rgba(0,0,0,.3);max-width:340px;font-size:13px}
  .hdr h1{font-size:15px;margin:0 0 4px}
  .lg{position:absolute;z-index:1000;bottom:14px;left:10px;background:#fff;padding:8px 10px;
      border-radius:6px;box-shadow:0 1px 6px rgba(0,0,0,.3);font-size:12px;line-height:1.7}
  .sw{display:inline-block;width:12px;height:12px;border-radius:50%;margin-right:6px;vertical-align:-1px}
  .ln{display:inline-block;width:14px;height:3px;margin-right:6px;vertical-align:3px}
</style></head><body>
<div id="map"></div>
<div class="hdr"><h1>Union Ave — struck far from a safe crossing</h1>
  Pedestrian/non-motorist crashes vs the combined safe-crossing inventory
  (__NSIG__ signalized + __NMK__ marked-only). Crash dots shaded by along-corridor distance to the
  nearest safe crossing; the longest no-crossing gap is highlighted.</div>
<div class="lg">
  <span class="ln" style="background:#e63946"></span>longest gap (no safe crossing)<br>
  <span class="ln" style="background:#1d8a4e"></span>signalized crosswalk (TDOT)<br>
  <span class="ln" style="background:#2a6fb0"></span>marked-only crosswalk (OSM)<br>
  <span class="sw" style="background:#1d8a4e"></span>TDOT signal head/button (per corner)<br>
  <span class="sw" style="background:#b30000"></span>fatal crash &nbsp;
  <span class="sw" style="background:#f4a259"></span>injury crash<br>
  <span style="color:#888">dot size = distance to nearest crossing</span>
</div>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script>
var map=L.map('map',{preferCanvas:true}).setView(__CENTER__,13);
L.tileLayer('https://{s}.basemaps.cartocdn.com/light_all/{z}/{x}/{y}{r}.png',
  {attribution:'&copy; OpenStreetMap, &copy; CARTO',maxZoom:20}).addTo(map);
L.polyline(__REF__,{color:'#888',weight:2,opacity:.7}).addTo(map);
L.polyline(__GAP__,{color:'#e63946',weight:7,opacity:.55}).addTo(map)
  .bindPopup('Longest stretch of Union with no safe crossing');
var COL={'signalized':'#1d8a4e','marked-only':'#2a6fb0'};
L.geoJSON(__OSM__,{style:function(f){return {color:COL[f.properties.type]||'#2a6fb0',weight:5,opacity:.85};},
  pointToLayer:function(f,ll){return L.circleMarker(ll,{radius:5,color:COL[f.properties.type],
    fillColor:COL[f.properties.type],fillOpacity:.9,weight:1});},
  onEachFeature:function(f,l){l.bindPopup((f.properties.type)+' crosswalk');}}).addTo(map);
L.geoJSON(__TDOT__,{pointToLayer:function(f,ll){return L.circleMarker(ll,{radius:2.5,color:'#1d8a4e',
  fillColor:'#1d8a4e',fillOpacity:.8,weight:0});},
  onEachFeature:function(f,l){l.bindPopup('TDOT '+(f.properties.FEATURE_DESCRIPTION||'signal'));}}).addTo(map);
function rad(ft){if(ft==null)return 5;return Math.max(5,Math.min(16,5+ft/40));}
L.geoJSON(__CRASH__,{pointToLayer:function(f,ll){
  var p=f.properties, fill=p.fatal?'#b30000':(p.nonrel?'#9aa':'#f4a259');
  return L.circleMarker(ll,{radius:p.fatal?rad(p.d_ft)+1:rad(p.d_ft),color:'#400',weight:p.fatal?1.5:.6,
    fillColor:fill,fillOpacity:p.fatal?.95:.8});},
  onEachFeature:function(f,l){var p=f.properties;
    l.bindPopup('<b>'+(p.fatal?'Fatal':'Injury')+'</b><br>'+p.date+'<br>'+p.loc+'<br>'+
      (p.nonrel?'<i>non-crossing</i>':('Nearest safe crossing: <b>'+(p.d_ft==null?'n/a':p.d_ft+' ft')+'</b>')));}}).addTo(map);
</script></body></html>"""


if __name__ == "__main__":
    main()
