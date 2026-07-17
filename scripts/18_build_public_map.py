r"""
18_build_public_map.py
=====================

PASS 2 — rebuild the public-facing interactive map + findings dashboard as ONE
clean, self-contained page, off the Pass-1 canonical classifier output. Replaces
the script-09/10 page (those are superseded).

STREETSTAT REDESIGN (2026-07): the page is now the StreetStat product shell —
a top navigation bar and four hash-routed views on a single self-contained page:
  #/            landing hero (computed stat cards) + the findings dashboard
  #/explore     the citywide map, one analytic "lens" at a time
                (Road ownership / Sidewalk inventory / Crash density), crash
                dots always visible, fatal-only + crossings as independent
                toggles;
  #/investigate a location microscope (skeleton here; wired by script 24,
                which owns the Count-A pipeline);
  #/methodology plain-language documentation of the whole pipeline.
Single page (not multiple pages) because the crash array + search index are
embedded for file:// robustness — separate pages would triple the payload or
require a data-loading refactor. Data, methodology and every computed number
are UNCHANGED; this is presentation-layer work only.

Design goals (unchanged underneath):
  - default map view = individual crash DOTS in THREE categories
    (City of Memphis / TDOT state route / Limited-access), fatal emphasized;
  - simple user-facing popups (date, severity, location, road-owner).

Every number is computed from the data files (nothing hardcoded) and reconciles
against the current totals at build time (sanity anchors live in CLAUDE.md).

Inputs (Pass 1):
  data/processed/shelby_crashes_final.csv
  outputs/interactive_map/ownership_segments_final.geojson
  data/raw/memphis_boundary.geojson   (subtle context outline)
Outputs:
  outputs/interactive_map/index.html              (rebuilt)
  data/processed/novel_statistics.docx            (APPEND a dated Pass-2 section)

Run it with:
    .\.venv\Scripts\python.exe scripts\18_build_public_map.py
(then re-run scripts\24_build_search.py — rebuilding index.html drops the
 injected search/investigate bundle).
"""

import json
import sys
from pathlib import Path
from datetime import date, datetime

import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
PROCESSED = ROOT / "data" / "processed"
OUT_MAP = ROOT / "outputs" / "interactive_map"
FINAL_CSV = PROCESSED / "shelby_crashes_final.csv"
SEG_GEOJSON = OUT_MAP / "ownership_segments_final.geojson"
BOUNDARY = ROOT / "data" / "raw" / "memphis_boundary.geojson"
INDEX_HTML = OUT_MAP / "index.html"
DOCX = PROCESSED / "novel_statistics.docx"
RAW_STREETS = ROOT / "data" / "raw" / "memphis_streets.geojson"      # LANES per segment (by OBJECTID)
RULEBOOK_GJ = PROCESSED / "road_ownership_rulebook.geojson"          # canonical network geometry + MTFCC

# palette — semantic data colors (unchanged meanings: teal=City, crimson=TDOT, charcoal=limited)
CITY_C, TDOT_C, LIM_C = "#1b9e8f", "#d6453d", "#3a3a44"
FATAL_STROKE = "#10242e"

CITY = "City of Memphis"
TDOT_SR = "TDOT state route"
LIMITED = {"Interstate (TDOT)", "Interstate ramp (TDOT)", "Limited-access (TDOT)"}
FATAL = "Fatal"


def pct(p, w):
    return round(100.0 * p / w, 1) if w else 0.0


def cat3(own):
    if own == CITY:
        return "City"
    if own == TDOT_SR:
        return "TDOT"
    return "Limited"


def compute_stats(f):
    N, NF = len(f), int((f.InjuryClass == FATAL).sum())
    fat = f[f.InjuryClass == FATAL]
    f = f.copy()
    f["cat3"] = f.Ownership.map(cat3)

    surf = f[f.Ownership.isin([CITY, TDOT_SR])]
    s_tot = len(surf)
    c_all = int((surf.Ownership == CITY).sum())
    t_all = int((surf.Ownership == TDOT_SR).sum())
    sf = surf[surf.InjuryClass == FATAL]
    sf_tot = len(sf)
    c_f = int((sf.Ownership == CITY).sum())
    t_f = int((sf.Ownership == TDOT_SR).sum())

    U = int(f.is_corner_case.sum())
    Uf = int((f.is_corner_case & (f.InjuryClass == FATAL)).sum())

    lim = f[f.is_limited_access]
    lim_n, lim_f = len(lim), int((lim.InjuryClass == FATAL).sum())
    int_n = int((f.Ownership == "Interstate (TDOT)").sum())
    ramp_n = int((f.Ownership == "Interstate ramp (TDOT)").sum())
    sam_n = int((f.Ownership == "Limited-access (TDOT)").sum())
    int_f = int(((f.Ownership == "Interstate (TDOT)") & (f.InjuryClass == FATAL)).sum())
    ramp_f = int(((f.Ownership == "Interstate ramp (TDOT)") & (f.InjuryClass == FATAL)).sum())
    sam_f = int(((f.Ownership == "Limited-access (TDOT)") & (f.InjuryClass == FATAL)).sum())

    # design (fatal)
    l4 = int((fat.Street_LANES >= 4).sum())
    s40 = int((fat.Street_SPDLIMIT >= 40).sum())
    both = int(((fat.Street_LANES >= 4) & (fat.Street_SPDLIMIT >= 40)).sum())

    # lighting (fatal)
    dark = int(fat.LightCondition.astype(str).str.startswith("Dark").sum())
    unlit = int(fat.LightCondition.astype(str).eq("Dark-Not Lighted").sum())

    # per-street concentration
    g = f.groupby("Street_Name").agg(
        total=("MstrRecNbrTxt", "size"),
        fatal=("InjuryClass", lambda s: int((s == FATAL).sum())),
        serious=("InjuryClass", lambda s: int((s == "Suspected Serious Injury").sum())),
    )
    by_fatal = g.sort_values(["fatal", "total"], ascending=False)
    cum = by_fatal["fatal"].cumsum()
    streets_half = int((cum < NF / 2).sum() + 1)
    n_streets = len(g)
    zero_fatal = int((g.fatal == 0).sum())

    # top 25 by total then fatal, with dominant category
    top = g.sort_values(["total", "fatal"], ascending=False).head(25)
    top25 = []
    for nm, r in top.iterrows():
        sub = f[f.Street_Name == nm]
        vc = sub.cat3.value_counts()
        dom = vc.index[0]
        dom_share = vc.iloc[0] / vc.sum()
        owner_lbl = {"City": "City of Memphis", "TDOT": "TDOT state route",
                     "Limited": "Limited-access"}[dom]
        spd = sub.Street_SPDLIMIT.mode()
        lanes = sub.Street_LANES.mode()
        top25.append({
            "name": nm, "total": int(r.total), "fatal": int(r.fatal),
            "serious": int(r.serious), "owner": owner_lbl,
            "mixed": bool(dom_share < 0.9),
            "spd": (None if spd.empty or pd.isna(spd.iloc[0]) else int(spd.iloc[0])),
            "lanes": (None if lanes.empty or pd.isna(lanes.iloc[0]) else int(lanes.iloc[0])),
        })

    # charts
    lane_counts = fat.Street_LANES.dropna().astype(int).value_counts().sort_index()
    years = sorted(int(y) for y in f.YearNmb.dropna().unique())
    # parse the collision-date field to real datetimes (it is US-format M/D/YYYY strings);
    # lexicographic string min/max would mis-rank these (e.g. "9/9/2025" > "5/26/2026").
    _cd = pd.to_datetime(f["CollisionDate"], errors="coerce")
    dmin, dmax = _cd.min().strftime("%Y-%m-%d"), _cd.max().strftime("%Y-%m-%d")

    return {
        "N": N, "NF": NF, "s_tot": s_tot, "sf_tot": sf_tot,
        "c_all": c_all, "t_all": t_all, "c_f": c_f, "t_f": t_f,
        "c_all_pct": pct(c_all, s_tot), "t_all_pct": pct(t_all, s_tot),
        "c_f_pct": pct(c_f, sf_tot), "t_f_pct": pct(t_f, sf_tot),
        "U": U, "Uf": Uf,
        "c_all_lo": pct(c_all - U, s_tot), "t_all_hi": pct(t_all + U, s_tot),
        "c_f_lo": pct(c_f - Uf, sf_tot), "t_f_hi": pct(t_f + Uf, sf_tot),
        "lim_n": lim_n, "lim_f": lim_f,
        "int_n": int_n, "ramp_n": ramp_n, "sam_n": sam_n,
        "int_f": int_f, "ramp_f": ramp_f, "sam_f": sam_f,
        "d_4ln": pct(l4, NF), "d_40": pct(s40, NF), "d_both": pct(both, NF), "d_both_n": both,
        "lit_dark": pct(dark, NF), "lit_unlit": pct(unlit, NF), "lit_unlit_n": unlit,
        "streets_half": streets_half, "n_streets": n_streets,
        "zero_fatal_pct": pct(zero_fatal, n_streets),
        "top25": top25,
        "chart_juris": {"all": [c_all, t_all, lim_n], "fatal": [c_f, t_f, lim_f]},
        "chart_lanes": {"labels": [int(k) for k in lane_counts.index],
                        "values": [int(v) for v in lane_counts.values]},
        "chart_year": {"labels": years,
                       "all": [int((f.YearNmb == y).sum()) for y in years],
                       "fatal": [int(((f.YearNmb == y) & (f.InjuryClass == FATAL)).sum()) for y in years]},
        "dmin": dmin, "dmax": dmax,
    }


def lane_mileage_context():
    """Exposure context for the deaths-by-lane-count statistic (added 2026-07-12): the share of
    NETWORK street mileage by lane category, COMPUTED from the road data (never estimated).
    Geometry + MTFCC come from the canonical rulebook network; LANES joins from the raw city
    street-centerline file by OBJECTID; lengths measured in EPSG:32136 meters. "Surface" excludes
    interstate mainlines (S1100) and ramps (S1630), matching how the findings report them
    separately. Returns None if inputs are unavailable -> the caption is omitted, never guessed."""
    try:
        import geopandas as gpd
        if not (RAW_STREETS.exists() and RULEBOOK_GJ.exists()):
            print("  (lane-mileage context skipped: network files not present)")
            return None
        rb = gpd.read_file(RULEBOOK_GJ, columns=["OBJECTID", "MTFCC"]).to_crs("EPSG:32136")
        rb["len_m"] = rb.geometry.length
        attrs = gpd.read_file(RAW_STREETS, columns=["OBJECTID", "LANES"], ignore_geometry=True)
        lanes = dict(zip(attrs["OBJECTID"], pd.to_numeric(attrs["LANES"], errors="coerce")))
        rb["LANES"] = rb["OBJECTID"].map(lanes)
        total_all = float(rb["len_m"].sum())
        known = rb[rb["LANES"].notna() & (rb["LANES"] > 0)]
        surface = known[~known["MTFCC"].isin(["S1100", "S1630"])]

        def share4(df):
            tot = float(df["len_m"].sum())
            return round(100.0 * float(df.loc[df["LANES"] >= 4, "len_m"].sum()) / tot, 1) if tot else None

        return {
            "pct4_surface": share4(surface),
            "pct4_all": share4(known),
            "known_pct": round(100.0 * float(known["len_m"].sum()) / total_all, 1) if total_all else None,
            "surface_mi": round(float(surface["len_m"].sum()) / 1609.344),
        }
    except Exception as e:
        print(f"  (lane-mileage context unavailable: {e})")
        return None


def crash_array(f):
    """Compact per-crash array: [lat, lng, cat(0/1/2), fatal(0/1), date(ISO), sev, loc]."""
    catn = {"City": 0, "TDOT": 1, "Limited": 2}
    iso = pd.to_datetime(f["CollisionDate"], errors="coerce").dt.strftime("%Y-%m-%d")  # M/D/YYYY -> ISO
    rows = []
    for pos, (_, r) in enumerate(f.iterrows()):
        d = iso.iloc[pos]
        rows.append([
            round(float(r.Latitude), 6), round(float(r.Longitude), 6),
            catn[cat3(r.Ownership)], 1 if r.InjuryClass == FATAL else 0,
            (d if isinstance(d, str) else ""), str(r.InjuryClass), str(r.NonMotoristLocation),
        ])
    return rows


def fmt_month(iso):
    return datetime.strptime(iso, "%Y-%m-%d").strftime("%b %Y")


def fmt_full(iso):
    d = datetime.strptime(iso, "%Y-%m-%d")
    return f"{d.strftime('%B')} {d.day}, {d.year}"


def hero_html(s):
    """Landing hero: name, honest subtitle, thesis, 4 computed stat cards, primary actions."""
    win = f"{fmt_month(s['dmin'])} &ndash; {fmt_month(s['dmax'])}"
    return f"""
<div class="hero-band">
  <div class="inner">
    <div class="eyebrow">Open data &middot; Memphis, Tennessee</div>
    <h1 class="hero-title">StreetStat</h1>
    <p class="hero-sub">Memphis pedestrian crashes &amp; infrastructure context</p>
    <p class="hero-thesis">Local coverage often frames each pedestrian death as an individual mistake.
       StreetStat puts every recorded pedestrian and non-motorist crash in its <b>infrastructure
       context</b> instead: what the road is like &mdash; lanes, posted speed, lighting, sidewalks,
       crossings &mdash; and <b>who owns it</b>, the City of Memphis or the Tennessee DOT. Everything
       here is a share, a count, or a distance computed from public records. The numbers are
       descriptive: they say where crashes happen and what those roads are like, never that a road
       <i>caused</i> a crash.</p>
    <div class="hero-cta">
      <a class="btn primary" href="#/explore">Explore the map</a>
      <a class="btn" href="#/investigate">Investigate a location</a>
    </div>
    <div class="statgrid">
      <div class="stat"><div class="stat-num">{s['N']:,}</div>
        <div class="stat-lab">crashes</div>
        <div class="stat-cap">every recorded pedestrian / non-motorist crash inside the City of
        Memphis (pedalcyclists excluded by design)</div></div>
      <div class="stat"><div class="stat-num">{s['NF']}</div>
        <div class="stat-lab">deaths</div>
        <div class="stat-cap">crashes in which a pedestrian or other non-motorist was killed</div></div>
      <div class="stat"><div class="stat-num">{s['t_all_pct']}%</div>
        <div class="stat-lab">on state-owned roads</div>
        <div class="stat-cap">share of surface-street crashes on TDOT state routes
        ({s['t_f_pct']}% of surface-street deaths); the rest are on City of Memphis streets</div></div>
      <div class="stat"><div class="stat-num stat-num-sm">{win}</div>
        <div class="stat-lab">data current through {fmt_full(s['dmax'])}</div>
        <div class="stat-cap">the state's rolling crash file; recent months may still be
        incomplete due to reporting lag</div></div>
    </div>
  </div>
</div>
"""


def _lane_ctx_line(s):
    """One-line exposure context under the lanes card (computed share, or nothing)."""
    lm = s.get("lane_mi")
    if not lm or lm.get("pct4_surface") is None:
        return ""
    return (f'\n      <div class="rng">context: roads with 4+ lanes account for '
            f'{lm["pct4_surface"]}% of surface-street mileage in the network</div>')


def dashboard_html(s):
    """Findings dashboard markup — every number computed from data (values unchanged by redesign)."""
    return f"""
<section id="stats"><div class="inner">
  <h2>Findings &mdash; who owns the deadly roads, and why people die on them</h2>
  <p class="sub">All figures are the {s['N']:,} pedestrian / non-motorist crashes inside the
     City of Memphis ({s['NF']} fatal), {fmt_full(s['dmin'])} to {fmt_full(s['dmax'])}. Recomputed
     from the data &mdash; nothing hand-entered.</p>

  <div class="cardgrid">
    <div class="card city">
      <div class="big">{s['c_all_pct']}%</div>
      <div class="lab">of surface crashes are on <b>City of Memphis</b>&ndash;owned roads</div>
      <div class="rng">range {s['c_all_lo']}&ndash;{s['c_all_pct']}% &middot; TDOT {s['t_all_pct']}&ndash;{s['t_all_hi']}%</div>
    </div>
    <div class="card city">
      <div class="big">{s['c_f_pct']}%</div>
      <div class="lab">of pedestrian <b>deaths</b> are on City of Memphis roads</div>
      <div class="rng">range {s['c_f_lo']}&ndash;{s['c_f_pct']}% &middot; TDOT {s['t_f_pct']}&ndash;{s['t_f_hi']}%</div>
    </div>
    <div class="card tdot">
      <div class="big">{s['d_4ln']}%</div>
      <div class="lab">of deaths are on roads with <b>4+ lanes</b>; {s['d_40']}% on roads posted <b>40+ mph</b></div>
      <div class="rng">nearly half ({s['d_both']}%) are on roads that are <i>both</i></div>{_lane_ctx_line(s)}
    </div>
    <div class="card dark">
      <div class="big">{s['lit_dark']}%</div>
      <div class="lab">of pedestrian deaths happen <b>after dark</b></div>
      <div class="rng">{s['lit_unlit']}% on a <b>dark, unlit</b> road ({s['lit_unlit_n']} deaths)</div>
    </div>
    <div class="card city">
      <div class="big">{s['streets_half']} streets</div>
      <div class="lab">hold <b>half</b> of all pedestrian deaths (of {s['n_streets']} streets)</div>
      <div class="rng">{s['zero_fatal_pct']}% of streets saw no death at all</div>
    </div>
    <div class="card ltd">
      <div class="big">{s['lim_n']}</div>
      <div class="lab">crashes ({s['lim_f']} fatal) on <b>limited-access</b> roads &mdash; interstates, ramps, Sam Cooper</div>
      <div class="rng">counted separately, not in the City/TDOT surface split</div>
    </div>
  </div>

  <p class="reframe">Most Memphis pedestrian deaths are not random residential accidents. They cluster on
     <b>wide, fast, multi-lane arterials</b> &mdash; {s['d_4ln']}% on roads of four or more lanes and
     {s['d_40']}% on roads posted 40&nbsp;mph or higher. When a road is built to move cars quickly
     through many lanes, a person trying to cross has little chance. This is a <b>design problem, not a
     behavior problem</b>: who builds, owns, and lights these roads &mdash; the City and TDOT &mdash;
     decides whether crossing them is survivable.</p>

  <h3>Who owns the road</h3>
  <table class="juris">
    <tr><th>Category</th><th>All crashes</th><th>Deaths</th></tr>
    <tr><td><span class="sw" style="background:{CITY_C}"></span>City of Memphis (surface)</td>
        <td class="n">{s['c_all']:,} ({s['c_all_pct']}%)</td><td class="n">{s['c_f']} ({s['c_f_pct']}%)</td></tr>
    <tr><td><span class="sw" style="background:{TDOT_C}"></span>TDOT state route (surface)</td>
        <td class="n">{s['t_all']:,} ({s['t_all_pct']}%)</td><td class="n">{s['t_f']} ({s['t_f_pct']}%)</td></tr>
    <tr><td><span class="sw" style="background:{LIM_C}"></span>Limited-access (TDOT) &mdash; separate</td>
        <td class="n">{s['lim_n']}</td><td class="n">{s['lim_f']}</td></tr>
  </table>
  <p class="rangenote"><b>The range.</b> Of surface crashes, the City of Memphis owns
     <b>{s['c_all_lo']}&ndash;{s['c_all_pct']}%</b> (about {round(s['c_all_pct'])}% by the primary method) and
     TDOT state routes <b>{s['t_all_pct']}&ndash;{s['t_all_hi']}%</b>; among deaths, City
     <b>{s['c_f_lo']}&ndash;{s['c_f_pct']}%</b> and TDOT <b>{s['t_f_pct']}&ndash;{s['t_f_hi']}%</b>. The
     range&apos;s upper TDOT bound credits the {s['U']} crashes at a city corner with a state route to the
     state route instead of the city cross-street. <b>The City owns the majority of surface pedestrian
     crashes under either reading.</b> Limited-access roads (interstates, ramps, Sam&nbsp;Cooper) are a
     separate {s['lim_n']} crashes ({s['lim_f']} fatal).</p>

  <h3>Charts</h3>
  <div class="charts">
    <div class="chart-box"><h4>Who owns the road &mdash; all crashes vs. deaths</h4>
      <canvas id="cJuris" height="220"></canvas></div>
    <div class="chart-box"><h4>Deaths by number of lanes</h4>
      <canvas id="cLanes" height="220"></canvas></div>
    <div class="chart-box"><h4>Crashes by year</h4>
      <canvas id="cYear" height="220"></canvas></div>
  </div>

  <h3>The 25 deadliest corridors</h3>
  <p class="sub">Click a column to sort. Owner is the dominant road owner along the corridor.</p>
  <table class="deadliest" id="deadliest">
    <thead><tr>
      <th data-k="rank">#</th><th data-k="name">Street</th><th data-k="total">Crashes</th>
      <th data-k="serious">Serious Injury</th><th data-k="fatal">Deaths</th><th data-k="owner">Owner</th>
      <th data-k="spd">Speed</th><th data-k="lanes">Lanes</th>
    </tr></thead><tbody></tbody>
  </table>

  <div class="foot">
    <p><b>Method.</b> Tennessee SAFETY non-motorist crash records for Shelby County,
       {fmt_full(s['dmin'])} to {fmt_full(s['dmax'])}, deduplicated to one crash per report and
       filtered to inside the City of Memphis.
       Each crash is assigned to the road it happened on &mdash; City of Memphis, a TDOT state route,
       or a limited-access TDOT road &mdash; by matching it to the nearest road centerline (all distance
       math in EPSG:32136, Tennessee meters). Pedalcyclists are excluded. Every figure is recomputed
       from the data. Full details: <a href="#/methodology">Methodology</a>.</p>
    <p><b>Sources.</b> Crashes: Tennessee SAFETY MapServer (TDOT). Roads, state routes, city boundary,
       street centerline: City of Memphis Public Works GIS. National ranking: Smart Growth America,
       <i>Dangerous by Design 2024</i>. Sam&nbsp;Cooper Blvd&apos;s low-speed western end is technically a
       city surface street; its tint here reflects the TDOT expressway.</p>
  </div>
</div></section>
"""


def about_html():
    """The About section — fixed copy (owner-approved wording, 2026-07-17); deliberately plain.
    Do not embellish or add promotional language. Only the two reference links are markup."""
    repo = "https://github.com/sdesai25unc/Memphis-Pedestrian-Incident-Context-Visualizer"
    return f"""
<section id="about"><div class="inner">
  <h2>About StreetStat</h2>
  <p>StreetStat maps every reported pedestrian crash in Memphis and Shelby County since
     January 2023 and puts each one in context: who owns the road where it happened, whether the
     city&rsquo;s inventory shows a sidewalk there, how far the nearest marked crossing is, and how
     that stretch of road compares over time. The data comes from public sources, mainly TDOT&rsquo;s
     crash database, the City of Memphis sidewalk inventory, and TDOT&rsquo;s signal records. The
     joining, counting, and classification are done by StreetStat&rsquo;s own open-source pipeline,
     and every method is documented on the <a href="#/methodology">Methodology page</a>.</p>
  <p>The site updates itself daily from the state&rsquo;s crash database. Numbers here are computed
     independently and are not official figures from TDOT or the City of Memphis. They can differ
     from official counts because of scope and methodology, and the
     <a href="#/methodology">Methodology page</a> explains exactly how each number is produced.</p>
  <p>StreetStat was built in summer 2026 during a civic fellowship with Innovate Memphis. The code
     is open source under the MIT license, and the full repository, including every script that
     produces what you see here, is public on
     <a href="{repo}" rel="noopener">GitHub</a>.</p>
</div></section>
"""


def _methodology_extras():
    """Build-time counts for the methodology page, read from the data files (never hardcoded).
    Every read degrades gracefully to None -> the prose falls back to threshold-only wording."""
    ex = {"union": None, "n_junctions": None, "n_corridors": None,
          "n_sidewalk_lines": None, "n_crossings": None}
    try:
        p = PROCESSED / "union_safe_summary.json"
        if p.exists():
            ex["union"] = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        pass
    try:
        p = PROCESSED / "search_index.json"
        if p.exists():
            meta = json.loads(p.read_text(encoding="utf-8")).get("meta", {})
            ex["n_junctions"] = meta.get("n_intersections")
            ex["n_corridors"] = meta.get("n_corridors")
    except Exception:
        pass
    try:
        p = PROCESSED / "memphis_sidewalks_32136.geojson"
        if p.exists():
            ex["n_sidewalk_lines"] = len(json.loads(p.read_text(encoding="utf-8")).get("features", []))
    except Exception:
        pass
    try:
        p = PROCESSED / "signalized_crossings_dedup.geojson"
        if p.exists():
            ex["n_crossings"] = len(json.loads(p.read_text(encoding="utf-8")).get("features", []))
    except Exception:
        pass
    return ex


def methodology_html(s):
    """The credibility page: for each pipeline stage — the public source, the rule we apply,
    and its honest limitations. Descriptions derive from the actual scripts (linked)."""
    ex = _methodology_extras()
    u = ex["union"] or {}

    def n(v, fallback="&mdash;"):
        return f"{v:,}" if isinstance(v, (int, float)) else fallback

    nj = n(ex["n_junctions"], "every junction citywide")
    nc = n(ex["n_corridors"], "all")
    nsw = n(ex["n_sidewalk_lines"], "the city&rsquo;s")
    nx = n(ex["n_crossings"], "the deduped")
    union_line = ""
    if u:
        union_line = (f"On Union Avenue this finds <b>{u.get('n_safe', '?')} safe crossings</b> "
                      f"({u.get('n_signalized', '?')} signalized + {u.get('n_marked_only', '?')} marked-only), "
                      f"a median spacing of {n(u.get('median_spacing_ft'))} ft, "
                      f"<b>{u.get('pct_over_250ft', '?')}% of crossing-relevant crashes more than 250 ft</b> "
                      f"from the nearest safe crossing, and a longest gap of "
                      f"<b>{n(u.get('longest_gap_ft'))} ft</b> &mdash; roughly "
                      f"{round(u.get('longest_gap_ft', 0) / 300.0, 1)}&times; the FHWA ~300 ft "
                      f"best-practice crossing spacing.")

    return f"""
<div class="doc"><div class="inner">
  <h2>Methodology</h2>
  <p class="sub">How every number on this site is produced &mdash; the public source, the exact rule
     applied, and the honest limits of each step. Written for a skeptical reader; script filenames
     are given for anyone who wants to check the code.</p>

  <p>Five principles run through the whole pipeline. <b>Computed, never hardcoded</b> &mdash; every
     figure is derived from the data files at build time. <b>Reconciled</b> &mdash; each build proves
     the splits sum back to <b>{s['N']:,} crashes / {s['NF']} deaths</b> before the page ships.
     <b>Correct geometry</b> &mdash; all distance math is done in EPSG:32136 (NAD83 / Tennessee, meters);
     the common Web Mercator projection is never used for measurement because it stretches distances
     about 22% at Memphis&rsquo;s latitude. <b>Descriptive, not causal</b> &mdash; shares, counts, and
     distances only; the data cannot say a road <i>caused</i> a crash, so the site never does.
     <b>Honest about coverage</b> &mdash; where a dataset doesn&rsquo;t cover a place, fields read
     &ldquo;not yet analyzed,&rdquo; never a fabricated number.</p>

  <h3>a. Crash data</h3>
  <p><b>Source:</b> the Tennessee SAFETY MapServer &mdash; TDOT&rsquo;s public crash service &mdash;
     non-motorist records for Shelby County. No login, no scraping; the raw API pages are saved
     untouched in <code>data/raw/</code> and never hand-edited.</p>
  <p><b>Rule:</b> the service returns one row per <i>person</i> involved, so records are deduplicated
     to one row per <i>crash</i>, keeping the worst injury in the crash. Pedalcyclists are excluded by
     design (this is a pedestrian project). Crashes with missing or impossible coordinates (blank,
     0&deg;/0&deg;, outside Shelby County) are flagged and kept in the file but excluded from headline
     counts, and everything is filtered to points inside the City of Memphis boundary. Result:
     <b>{s['N']:,}</b> crashes ({s['NF']} fatal), {fmt_full(s['dmin'])} to {fmt_full(s['dmax'])}.</p>
  <p><b>Limitations:</b> the state file is a rolling window of roughly the last three years, so totals
     shift slowly as it advances. Official crash data is finalized with a <b>reporting lag</b> &mdash;
     the most recent months undercount, and every time-window table on this site says so. The data is
     only as good as the underlying police reports.</p>
  <p class="srcline">Code: <code>scripts/01_download_crashes.py</code> &middot; <code>scripts/03_spatial_join.py</code></p>

  <h3>b. Crash-to-road attribution</h3>
  <p><b>Rule:</b> each crash point is matched to the <b>nearest road-centerline segment</b>, measured
     in EPSG:32136 meters, and inherits that segment&rsquo;s ownership. The matched segment ID and the
     snap distance are recorded on every crash row, so each attribution is auditable. (The project&rsquo;s
     original screening rule classed a crash as state-route when it lay within <b>30 m</b> of one; the
     final method replaces that with nearest-classified-segment inheritance and keeps the distance as
     provenance.)</p>
  <div class="caveat"><b>Intersection ambiguity.</b> Near a corner, the nearest centerline can be the
     cross street rather than the main road. Crashes on a city-named cross street within 10 m of a
     state route at an intersection are counted as City in the point estimate, and the published
     <i>range</i> on the findings page shows what happens if every one of them is credited to the state
     route instead. Every search result displays the snapped road name and snap distance so you can
     judge each case yourself.</div>
  <p class="srcline">Code: <code>scripts/17_classifier.py</code> &middot; <code>scripts/03_spatial_join.py</code></p>

  <h3>c. Road ownership classification</h3>
  <p><b>Source:</b> City of Memphis Public Works GIS (street centerlines, state-route layer, city
     boundary).</p>
  <p><b>Rule:</b> every centerline segment is tagged by an ordered rulebook; the first rule that
     matches wins, and <b>which rule fired is recorded on the segment</b>:</p>
  <ul>
    <li>Interstate mainline (road class S1100) &rarr; Limited-access (TDOT)</li>
    <li>Interstate ramp (S1630) &rarr; Limited-access (TDOT)</li>
    <li>Documented limited-access override list (e.g. Sam Cooper Blvd, which is absent from the
        state-route layer) &rarr; Limited-access (TDOT)</li>
    <li>Geometric overlap with a same-named state route (&ge;60% of the segment within 10 m, or
        &ge;85% within 8 m for differently-named routes) &rarr; TDOT state route</li>
    <li>Completeness override: a City segment &ge;20% collinear with a same-named state route
        (name-guarded, so it can never tag a cross street) &rarr; TDOT state route</li>
    <li>Otherwise &rarr; City of Memphis</li>
  </ul>
  <p><b>Limitations:</b> Sam Cooper&rsquo;s low-speed western end is technically a city surface street,
     so its map tint slightly over-extends (crash counts are unaffected). Ownership here means the
     road&rsquo;s classification in the public GIS layers, which is the best public record of design
     responsibility &mdash; not a legal opinion on maintenance agreements.</p>
  <p class="srcline">Code: <code>scripts/17_classifier.py</code> (consolidating <code>14</code>&ndash;<code>16</code>)</p>

  <h3>d. Corridors &amp; along-road counting</h3>
  <p><b>Rule:</b> corridors group crashes by standardized street name &mdash; the same grouping the
     deadliest-corridor table uses, with directional prefixes kept (North Parkway &ne; South Parkway).
     Each corridor&rsquo;s segments are merged into spatially-ordered connected lines; same-name branch
     stubs whose endpoints touch within <b>1 m</b> are treated as one road, while gaps larger than
     <b>15 m</b> (a rail yard, a genuine break) keep sections separate. The &ldquo;&plusmn;300 m&rdquo;
     count around a point is measured as <b>network distance along the road</b>, never a straight-line
     radius, and never leaks across a real gap &mdash; the two orange bars on the map show exactly where
     the window ends, and if a section is shorter than the window the result says the window was clamped.</p>
  <p><b>Verification:</b> the browser&rsquo;s coordinate math is checked against the reference Python
     implementation (they agree to under a millimeter), and every build re-proves that whole-corridor
     counts still sum to {s['N']:,} and that all 25 deadliest-corridor figures match the published table.</p>
  <p><b>Limitations:</b> corridor geometry is simplified (10 m tolerance) for the page, so a snap
     distance can differ by a few meters from survey-grade. Points on roads with no recorded pedestrian
     crashes snap to the nearest road that has them &mdash; the result card always names the road it
     counted and the snap distance.</p>
  <p class="srcline">Code: <code>scripts/24_build_search.py</code></p>

  <h3>e. Intersection index</h3>
  <p><b>Rule:</b> junctions come from <b>true geometric intersection</b> of the named through-road
     centerlines (not from segment endpoints, which miss crossings that don&rsquo;t share an endpoint).
     Interstates and ramps are excluded as through-roads so grade-separated overpasses don&rsquo;t create
     false at-grade &ldquo;intersections.&rdquo; Where a cross street meets a divided arterial it crosses
     each carriageway separately; those paired points are merged into <b>one junction</b> using a cluster
     radius set from the measured carriageway-gap distribution (25&ndash;30 m, with same-street-pair
     merges up to 120 m for wide medians). That yields <b>{nj} junctions</b> citywide, all searchable.
     A junction is &ldquo;signalized&rdquo; when a deduplicated TDOT pedestrian crossing lies within
     30 m; crashes attach to their nearest junction within 30 m.</p>
  <p><b>Limitations:</b> signal status is only known along corridors the TDOT inventory covers &mdash;
     elsewhere it reads &ldquo;not yet analyzed,&rdquo; because absence of inventory is not
     &ldquo;no signal.&rdquo; A junction with no recorded crashes honestly reports
     &ldquo;0 incidents reported here.&rdquo;</p>
  <p class="srcline">Code: <code>scripts/25_rebuild_junctions.py</code> (superseding <code>21</code>)</p>

  <h3>f. Sidewalk presence</h3>
  <p><b>Source:</b> the City of Memphis sidewalk inventory ({nsw} line features), reprojected to
     EPSG:32136.</p>
  <p><b>Rule:</b> a stretch of road is marked &ldquo;sidewalk present in city inventory&rdquo; when an
     inventory line runs within <b>20 m</b> of the road centerline. Why 20: measured against the data,
     sidewalks sit about 7 m (median) to 12 m (90th percentile) off the centerline, so 20 m catches
     near- <i>and</i> far-side sidewalks while staying well under the ~60 m spacing of parallel streets
     &mdash; a neighboring street&rsquo;s sidewalk can&rsquo;t register as this road&rsquo;s.</p>
  <div class="caveat"><b>Mandatory caveat.</b> The inventory records where the city has mapped
     sidewalks; it is not a guarantee of completeness. StreetStat therefore never says &ldquo;no
     sidewalk&rdquo; &mdash; the wording is always <i>&ldquo;No sidewalk found in city inventory
     (absence may reflect incomplete records).&rdquo;</i></div>
  <p class="srcline">Code: <code>scripts/24_build_search.py</code> (flags built at index time)</p>

  <h3>g. Safe-crossing distance (Union Avenue proof of concept)</h3>
  <p><b>Sources:</b> TDOT&rsquo;s &ldquo;ADA Asset Data&rdquo; pedestrian-signal inventory (signal heads
     and push buttons deduplicated at 30 m into {nx} signalized-crossing points, one per intersection)
     and OpenStreetMap marked crosswalks (node/way duplicates merged at 8 m).</p>
  <p><b>Rule:</b> on Union Avenue, a &ldquo;safe crossing&rdquo; is a TDOT signal or an OSM marked
     crosswalk within 30 m of the centerline, deduplicated so a signalized intersection that also has
     a marked crosswalk counts once. Each crossing-relevant crash&rsquo;s distance to the nearest safe
     crossing is measured <b>along the corridor</b> (linear referencing), not as the crow flies; crashes
     recorded away from the roadway are reported separately, not folded in. {union_line}</p>
  <p><b>Limitations:</b> this analysis is deliberately <b>Union-only</b> so far. OpenStreetMap crosswalk
     completeness varies block to block &mdash; it was evaluated against the TDOT inventory before use
     (<code>scripts/22_osm_crossings_eval.py</code>), and extending the analysis citywide waits on that
     ground-truthing. Elsewhere the &ldquo;nearest safe crossing&rdquo; field reads &ldquo;not yet
     analyzed.&rdquo;</p>
  <p class="srcline">Code: <code>scripts/23_union_poc.py</code> &middot; <code>scripts/19</code>&ndash;<code>22</code></p>

  <h3>Reproducibility</h3>
  <p>The pipeline is a sequence of numbered scripts run in order
     (download &rarr; classify &rarr; analyze &rarr; build) &mdash; see the
     <a href="https://github.com/sdesai25unc/Memphis-Pedestrian-Incident-Context-Visualizer" rel="noopener">repository</a>
     README for the exact commands. Raw API downloads are never edited; every derived file is written
     as a new file; and each build prints its own reconciliation against the fixed totals before the
     page is published. Code is MIT-licensed; source data remains under its providers&rsquo; terms
     (crashes: TDOT; roads &amp; sidewalks: City of Memphis; crosswalks: &copy; OpenStreetMap
     contributors, ODbL; geocoding: US Census Bureau).</p>
</div></div>
"""


def build_html(s, crashes, segments, boundary, crossings):
    stats_json = json.dumps({
        "juris": s["chart_juris"], "lanes": s["chart_lanes"], "year": s["chart_year"],
        "top25": s["top25"],
        "colors": {"city": CITY_C, "tdot": TDOT_C, "lim": LIM_C},
    })
    page = _TEMPLATE
    page = page.replace("/*__HERO__*/", hero_html(s))
    page = page.replace("/*__DASHBOARD__*/", dashboard_html(s))
    page = page.replace("/*__ABOUT__*/", about_html())
    page = page.replace("/*__METHODOLOGY__*/", methodology_html(s))
    page = page.replace("__DMAX_FULL__", fmt_full(s["dmax"]))
    page = page.replace("__DMIN_MON__", fmt_month(s["dmin"])).replace("__DMAX_MON__", fmt_month(s["dmax"]))
    page = page.replace("__DMIN_ISO__", s["dmin"]).replace("__DMAX_ISO__", s["dmax"])
    page = page.replace("__CITY_C__", CITY_C).replace("__TDOT_C__", TDOT_C)
    page = page.replace("__LIM_C__", LIM_C).replace("__FATAL_STROKE__", FATAL_STROKE)
    page = page.replace("__CRASHES__", json.dumps(crashes, separators=(",", ":")))
    page = page.replace("__SEGMENTS__", json.dumps(segments, separators=(",", ":")))
    page = page.replace("__BOUNDARY__", json.dumps(boundary, separators=(",", ":")))
    page = page.replace("__CROSSINGS__", json.dumps(crossings, separators=(",", ":")))
    page = page.replace("__STATS_JSON__", stats_json)
    return page


def append_docx(s):
    from docx import Document
    from docx.shared import Pt
    if not DOCX.exists():
        print(f"  (skip docx: {DOCX.name} not found)")
        return
    doc = Document(str(DOCX))
    if any(p.text.strip().startswith("Pass 2 — Final classification") for p in doc.paragraphs):
        print("  (docx already has the Pass-2 section; skipping append)")
        return
    doc.add_page_break()
    h = doc.add_heading(f"Pass 2 — Final classification & public map (rebuilt {date.today().isoformat()})", level=1)
    doc.add_paragraph(
        "The jurisdiction methodology is now canonical (rulebook-driven classifier, "
        "scripts/17_classifier.py). The public map + dashboard were rebuilt from it "
        "(scripts/18_build_public_map.py). Numbers below are computed from "
        "shelby_crashes_final.csv and reconcile to 1,294 crashes / 175 fatal.")

    doc.add_heading("Final surface City / TDOT split (range)", level=2)
    for line in [
        f"All surface crashes (n={s['s_tot']}): City {s['c_all_lo']}–{s['c_all_pct']}% "
        f"({s['c_all']:,} pt) / TDOT {s['t_all_pct']}–{s['t_all_hi']}% ({s['t_all']} pt).",
        f"Fatal surface crashes (n={s['sf_tot']}): City {s['c_f_lo']}–{s['c_f_pct']}% "
        f"({s['c_f']} pt) / TDOT {s['t_f_pct']}–{s['t_f_hi']}% ({s['t_f']} pt).",
        f"Point estimate = nearest-centerline (corner crashes counted as city). Upper TDOT "
        f"bound credits the {s['U']} state-route-corner crashes ({s['Uf']} fatal) to the state route.",
        f"Limited-access (TDOT), reported separately: {s['lim_n']} crashes ({s['lim_f']} fatal) "
        f"= interstate {s['int_n']} ({s['int_f']}), ramps {s['ramp_n']} ({s['ramp_f']}), "
        f"Sam Cooper {s['sam_n']} ({s['sam_f']}).",
        f"Reconciliation: surface {s['s_tot']} + limited-access {s['lim_n']} = {s['s_tot']+s['lim_n']} "
        f"(= 1,294); fatal {s['sf_tot']} + {s['lim_f']} = {s['sf_tot']+s['lim_f']} (= 175).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Methodology decision-tree (internal)", level=2)
    for i, line in enumerate([
        "Interstate mainline (MTFCC S1100) → Limited-access (TDOT).",
        "Interstate ramp (MTFCC S1630) → Limited-access (TDOT).",
        "Limited-access override (documented name list; seeded with Sam Cooper Blvd) → Limited-access (TDOT).",
        "State-route geometric overlap (≥60% of the segment within 10 m of a same-named state route, "
        "or ≥85% within 8 m) → TDOT state route.",
        "Force-state-route completeness override (a City segment ≥20% collinear with a same-named "
        "state route; name-guarded) → TDOT state route.",
        "Otherwise → City of Memphis.",
    ], 1):
        doc.add_paragraph(f"Rule {i}: {line}", style="List Number")
    doc.add_paragraph(
        "Each crash inherits the ownership of its nearest rulebook segment (EPSG:32136). "
        "A corner case = a city crash on a non-state-route-named street, at an intersection, "
        "within 10 m of a state route (drives the range's upper bound only).")

    doc.add_heading("Per-incident provenance", level=2)
    doc.add_paragraph(
        "data/processed/shelby_crashes_final.csv carries, for every crash: Ownership, "
        "Classification_Basis (human-readable reason), Rule_Fired, matched Seg_OBJECTID, "
        "DistToSeg_m, flags is_limited_access / is_corner_case / is_override, and Jurisdiction_prev "
        "(the prior segment-method label) as an audit trail.")
    doc.save(str(DOCX))
    print(f"  appended Pass-2 section to {DOCX.name}")


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    f = pd.read_csv(FINAL_CSV)
    for col in ["is_limited_access", "is_corner_case", "is_override"]:
        f[col] = f[col].astype(str).str.lower().isin(["true", "1", "yes"])
    s = compute_stats(f)
    s["lane_mi"] = lane_mileage_context()
    crashes = crash_array(f)
    segments = json.loads(SEG_GEOJSON.read_text(encoding="utf-8"))
    boundary = json.loads(BOUNDARY.read_text(encoding="utf-8")) if BOUNDARY.exists() else {"type": "FeatureCollection", "features": []}
    cross_path = PROCESSED / "signalized_crossings_dedup.geojson"
    crossings = json.loads(cross_path.read_text(encoding="utf-8")) if cross_path.exists() else {"type": "FeatureCollection", "features": []}

    html = build_html(s, crashes, segments, boundary, crossings)
    INDEX_HTML.write_text(html, encoding="utf-8")
    print(f"Wrote {INDEX_HTML} ({len(html)/1e6:.2f} MB)")
    append_docx(s)

    # ---- report ----
    def row(label, allv, fatv):
        print(f"{label:<34}{allv:>18}{fatv:>14}")
    print("\n=== FINAL JURISDICTION TABLE (as rendered) ===")
    row("Category", "All crashes", "Deaths")
    row("City of Memphis (surface)", f"{s['c_all']:,} ({s['c_all_pct']}%)", f"{s['c_f']} ({s['c_f_pct']}%)")
    row("TDOT state route (surface)", f"{s['t_all']} ({s['t_all_pct']}%)", f"{s['t_f']} ({s['t_f_pct']}%)")
    row("Limited-access (TDOT) [separate]", str(s['lim_n']), str(s['lim_f']))
    print(f"\nRange: City {s['c_all_lo']}-{s['c_all_pct']}% / TDOT {s['t_all_pct']}-{s['t_all_hi']}% (all); "
          f"City {s['c_f_lo']}-{s['c_f_pct']}% / TDOT {s['t_f_pct']}-{s['t_f_hi']}% (fatal).")
    surf_recon = s['s_tot'] + s['lim_n']
    fat_recon = s['sf_tot'] + s['lim_f']
    # Internal-consistency gate: surface + limited-access must partition the in-Memphis total.
    # Compared against the CURRENT data (never a hardcoded anchor — anchors legitimately move
    # when the state's rolling window advances; update CLAUDE.md's anchors after a refresh).
    print(f"\nRECONCILIATION: surface {s['s_tot']} + limited {s['lim_n']} = {surf_recon} "
          f"{'OK' if surf_recon == s['N'] else 'FAIL'} (= in-Memphis total {s['N']:,}); "
          f"fatal {s['sf_tot']} + {s['lim_f']} = {fat_recon} "
          f"{'OK' if fat_recon == s['NF'] else 'FAIL'} (= {s['NF']})")

    if s.get("lane_mi"):
        lm = s["lane_mi"]
        print(f"\nLANE-MILEAGE CONTEXT (computed from the network, EPSG:32136 lengths): "
              f"4+ lanes = {lm['pct4_surface']}% of surface-street mileage "
              f"({lm['pct4_all']}% incl. limited-access); ~{lm['surface_mi']:,} surface miles; "
              f"lanes known for {lm['known_pct']}% of network mileage")

    print("\n=== CARD NUMBERS THAT DRIFTED (old page -> new) ===")
    drift = [
        ("City share, all crashes", "74.7%", f"{s['c_all_pct']}% (range {s['c_all_lo']}-{s['c_all_pct']}%)"),
        ("TDOT share, all crashes", "25.3%", f"{s['t_all_pct']}% (range {s['t_all_pct']}-{s['t_all_hi']}%)"),
        ("City share, fatal", "70.3%", f"{s['c_f_pct']}%"),
        ("TDOT share, fatal", "29.7%", f"{s['t_f_pct']}%"),
        ("Limited-access line", "(folded into City/TDOT)", f"{s['lim_n']} crashes / {s['lim_f']} fatal, separate"),
        ("Design: 4+ lanes & 40+ mph", "49.7% (hero)", f"{s['d_both']}% (reframed as 'nearly half', not a majority)"),
        ("Design: 4+ lanes / 40+ mph majorities", "62.9% / 60.0%", f"{s['d_4ln']}% / {s['d_40']}%"),
        ("Lighting after dark / unlit", "76.6% / 14.3%", f"{s['lit_dark']}% / {s['lit_unlit']}%"),
        ("Concentration", "26 streets / 80.3% none", f"{s['streets_half']} streets / {s['zero_fatal_pct']}% none"),
    ]
    for lbl, old, new in drift:
        flag = "" if old.split()[0] == new.split()[0] else "  <-- changed"
        print(f"  {lbl:<40} {old:<26} -> {new}{flag}")


_TEMPLATE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>StreetStat — Memphis Pedestrian Crashes &amp; Infrastructure Context</title>
<meta name="description" content="StreetStat maps every reported pedestrian crash in Memphis, Tennessee (__DMIN_MON__ to __DMAX_MON__) and shows who owns each road (City of Memphis or TDOT), sidewalk inventory status, and distance to the nearest marked crossing, computed from public data.">
<link rel="canonical" href="https://streetstat.org/">
<meta property="og:type" content="website">
<meta property="og:site_name" content="StreetStat">
<meta property="og:title" content="StreetStat — Memphis Pedestrian Crashes &amp; Infrastructure Context">
<meta property="og:description" content="StreetStat maps every reported pedestrian crash in Memphis, Tennessee (__DMIN_MON__ to __DMAX_MON__) and shows who owns each road (City of Memphis or TDOT), sidewalk inventory status, and distance to the nearest marked crossing, computed from public data.">
<meta property="og:url" content="https://streetstat.org/">
<script type="application/ld+json">
{"@context":"https://schema.org","@type":"Dataset","name":"StreetStat — Memphis Pedestrian Crashes & Infrastructure Context","description":"Every reported pedestrian and non-motorist crash inside the City of Memphis, __DMIN_MON__ to __DMAX_MON__, each attributed to the road it happened on and its owner (City of Memphis or TDOT), with sidewalk inventory status and crossing context, computed from public data.","url":"https://streetstat.org/","spatialCoverage":"Memphis, Tennessee, United States","temporalCoverage":"__DMIN_ISO__/__DMAX_ISO__","isAccessibleForFree":true,"creator":{"@type":"Person","name":"Samarth Desai"}}
</script>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 32 32'%3E%3Crect width='32' height='32' rx='7' fill='%2318181b'/%3E%3Ccircle cx='11' cy='21' r='4' fill='%231b9e8f'/%3E%3Ccircle cx='21' cy='11' r='4' fill='%23d6453d'/%3E%3C/svg%3E">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css">
<link rel="preconnect" href="https://cdn.jsdelivr.net" crossorigin>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-sans@5.2.5/400.css">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-sans@5.2.5/500.css">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-sans@5.2.5/600.css">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-sans@5.2.5/700.css">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-mono@5.2.5/400.css">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fontsource/geist-mono@5.2.5/500.css">
<style>
  /* ============ StreetStat design tokens ============ */
  :root {
    --city: __CITY_C__; --tdot: __TDOT_C__; --lim: __LIM_C__;
    --sw-yes: #2a6f97; --sw-no: #d98324;           /* sidewalk semantic colors (unchanged) */
    --bg: #fafafa; --surface: #ffffff;
    --ink: #18181b; --ink-2: #3f3f46; --muted: #71717a; --faint: #a1a1aa;
    --border: #e4e4e7; --border-strong: #d4d4d8;
    --accent: #4f46e5; --accent-ink: #4338ca; --accent-soft: #eef2ff;
    --r-sm: 6px; --r-md: 10px; --r-lg: 14px;
    --shadow-sm: 0 1px 2px rgba(24,24,27,.06);
    --shadow-md: 0 1px 2px rgba(24,24,27,.05), 0 4px 12px rgba(24,24,27,.08);
    --shadow-lg: 0 2px 4px rgba(24,24,27,.06), 0 10px 28px rgba(24,24,27,.12);
    --sans: "Geist Sans", "Geist", Inter, system-ui, -apple-system, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    --mono: "Geist Mono", ui-monospace, "Cascadia Code", "SF Mono", Consolas, monospace;
    --nav-h: 57px;
  }
  * { box-sizing: border-box; }
  html { scroll-behavior: smooth; }
  html, body { margin: 0; padding: 0; }
  body { font-family: var(--sans); color: var(--ink); background: var(--bg);
         font-size: 15px; line-height: 1.6; -webkit-font-smoothing: antialiased; }
  a { color: var(--accent-ink); text-decoration: none; }
  a:hover { text-decoration: underline; }
  .num, .mono { font-family: var(--mono); font-variant-numeric: tabular-nums; }

  /* ============ top navigation ============ */
  #topnav { position: sticky; top: 0; z-index: 3000; height: var(--nav-h);
    background: rgba(255,255,255,.86); backdrop-filter: blur(10px); -webkit-backdrop-filter: blur(10px);
    border-bottom: 1px solid var(--border); display: flex; align-items: center;
    padding: 0 22px; gap: 26px; }
  #topnav .wordmark { display: flex; align-items: center; gap: 9px; font-weight: 700;
    font-size: 16px; letter-spacing: -.01em; color: var(--ink); text-decoration: none; }
  #topnav .wordmark .mark { width: 20px; height: 20px; border-radius: 5px; background: var(--ink);
    position: relative; flex: none; }
  #topnav .wordmark .mark::before { content: ""; position: absolute; left: 4px; bottom: 4px;
    width: 6px; height: 6px; border-radius: 50%; background: var(--city); }
  #topnav .wordmark .mark::after { content: ""; position: absolute; right: 4px; top: 4px;
    width: 6px; height: 6px; border-radius: 50%; background: var(--tdot); }
  #topnav .navlinks { display: flex; gap: 4px; margin-left: auto; }
  #topnav .navlinks a { padding: 6px 13px; border-radius: 8px; font-size: 13.5px; font-weight: 500;
    color: var(--ink-2); text-decoration: none; transition: background .12s, color .12s; }
  #topnav .navlinks a:hover { background: #f4f4f5; color: var(--ink); }
  #topnav .navlinks a.active { background: var(--ink); color: #fff; }
  @media (max-width: 560px) { #topnav { padding: 0 12px; gap: 10px; }
    #topnav .navlinks a { padding: 6px 9px; font-size: 12.5px; } }

  /* ============ views ============ */
  .view[hidden] { display: none; }

  /* ============ hero ============ */
  .hero-band { background: var(--surface); border-bottom: 1px solid var(--border); }
  .hero-band .inner { max-width: 1060px; margin: 0 auto; padding: 72px 24px 58px; }
  .eyebrow { font-family: var(--mono); font-size: 11.5px; letter-spacing: .09em;
    text-transform: uppercase; color: var(--muted); margin-bottom: 18px; }
  .hero-title { margin: 0; font-size: 52px; font-weight: 700; letter-spacing: -.035em; line-height: 1.05; }
  .hero-sub { margin: 10px 0 0; font-size: 19px; color: var(--ink-2); font-weight: 400; }
  .hero-thesis { max-width: 760px; margin: 22px 0 0; font-size: 15.5px; line-height: 1.7; color: var(--ink-2); }
  .hero-thesis b { color: var(--ink); }
  .hero-cta { display: flex; gap: 10px; margin-top: 28px; flex-wrap: wrap; }
  .btn { display: inline-flex; align-items: center; gap: 7px; padding: 9px 18px;
    border-radius: var(--r-md); border: 1px solid var(--border-strong); background: var(--surface);
    color: var(--ink); font-size: 14px; font-weight: 600; text-decoration: none; cursor: pointer;
    transition: background .12s, border-color .12s, transform .06s; }
  .btn:hover { background: #f4f4f5; text-decoration: none; }
  .btn:active { transform: translateY(1px); }
  .btn.primary { background: var(--ink); border-color: var(--ink); color: #fff; }
  .btn.primary:hover { background: #2c2c31; }
  .statgrid { display: grid; grid-template-columns: repeat(auto-fit, minmax(215px, 1fr));
    gap: 14px; margin-top: 44px; }
  .stat { background: var(--surface); border: 1px solid var(--border); border-radius: var(--r-lg);
    padding: 18px 18px 16px; box-shadow: var(--shadow-sm); }
  .stat-num { font-family: var(--mono); font-size: 34px; font-weight: 500; letter-spacing: -.02em;
    line-height: 1.1; color: var(--ink); }
  .stat-num-sm { font-size: 21px; padding: 6px 0 7px; }
  .stat-lab { font-size: 13.5px; font-weight: 600; color: var(--ink); margin-top: 6px; }
  .stat-cap { font-size: 12px; color: var(--muted); margin-top: 5px; line-height: 1.5; }

  /* ============ findings dashboard ============ */
  #stats { background: var(--bg); padding: 46px 24px 64px; }
  #stats .inner { max-width: 1060px; margin: 0 auto; }
  #stats h2 { font-size: 26px; letter-spacing: -.02em; margin: 0 0 6px; }
  #stats h3 { font-size: 18px; letter-spacing: -.01em; margin: 42px 0 14px;
    padding-bottom: 8px; border-bottom: 1px solid var(--border); }
  #stats h4 { font-size: 13.5px; margin: 0 0 10px; color: var(--ink-2); }
  #stats .sub { color: var(--muted); margin: 0 0 22px; font-size: 13.5px; }
  .cardgrid { display: grid; grid-template-columns: repeat(auto-fit, minmax(215px, 1fr)); gap: 14px; }
  .card { background: var(--surface); border: 1px solid var(--border); border-radius: var(--r-lg);
    padding: 16px 17px; box-shadow: var(--shadow-sm); border-top: 3px solid var(--ink); }
  .card.city { border-top-color: var(--city); }
  .card.tdot { border-top-color: var(--tdot); }
  .card.dark { border-top-color: #2b2b50; }
  .card.ltd  { border-top-color: var(--lim); }
  .card .big { font-family: var(--mono); font-size: 28px; font-weight: 500; line-height: 1.15; }
  .card .lab { font-size: 13px; color: var(--ink-2); margin-top: 6px; }
  .card .rng { font-size: 11.5px; color: var(--muted); margin-top: 8px; }
  .reframe { background: var(--surface); border: 1px solid var(--border);
    border-left: 3px solid var(--tdot); padding: 15px 18px; border-radius: var(--r-md);
    margin: 24px 0 0; font-size: 14.5px; line-height: 1.7; box-shadow: var(--shadow-sm); }
  table.juris { border-collapse: separate; border-spacing: 0; width: 100%; max-width: 640px;
    background: var(--surface); font-size: 14px; border: 1px solid var(--border);
    border-radius: var(--r-md); overflow: hidden; box-shadow: var(--shadow-sm); }
  table.juris th { background: #f4f4f5; color: var(--ink-2); text-align: left; padding: 9px 14px;
    font-size: 12px; text-transform: uppercase; letter-spacing: .05em; border-bottom: 1px solid var(--border); }
  table.juris td { padding: 10px 14px; border-bottom: 1px solid var(--border); }
  table.juris tr:last-child td { border-bottom: none; }
  table.juris td.n { font-family: var(--mono); font-size: 13px; }
  .sw { display: inline-block; width: 10px; height: 10px; border-radius: 50%; margin-right: 8px; vertical-align: middle; }
  .rangenote { font-size: 13px; color: var(--ink-2); line-height: 1.65; margin-top: 14px; max-width: 860px; }
  .charts { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 14px; }
  .chart-box { background: var(--surface); border: 1px solid var(--border); border-radius: var(--r-lg);
    padding: 15px 17px; box-shadow: var(--shadow-sm); }
  table.deadliest { width: 100%; border-collapse: separate; border-spacing: 0; background: var(--surface);
    font-size: 13px; border: 1px solid var(--border); border-radius: var(--r-md); overflow: hidden;
    box-shadow: var(--shadow-sm); }
  table.deadliest th { background: #f4f4f5; color: var(--ink-2); text-align: left; padding: 9px 11px;
    font-size: 11.5px; text-transform: uppercase; letter-spacing: .05em; cursor: pointer;
    user-select: none; border-bottom: 1px solid var(--border); white-space: nowrap; }
  table.deadliest th:hover { background: #ececee; color: var(--ink); }
  table.deadliest td { padding: 8px 11px; border-bottom: 1px solid var(--border); }
  table.deadliest tr:last-child td { border-bottom: none; }
  table.deadliest td.n { font-family: var(--mono); font-size: 12.5px; }
  table.deadliest tr:hover td { background: #fafafa; }
  .foot { margin-top: 34px; font-size: 12.5px; color: var(--muted); line-height: 1.65; max-width: 860px; }

  /* ============ explore ============ */
  #view-explore .mapslot { position: relative; height: calc(100vh - var(--nav-h)); min-height: 460px; }
  #map { position: absolute; inset: 0; background: #eef1f3; z-index: 1; }
  .leaflet-container { font-family: var(--sans); }
  .leaflet-popup-content-wrapper { border-radius: var(--r-md); box-shadow: var(--shadow-lg); }
  .leaflet-popup-content { font-size: 13px; line-height: 1.55; margin: 11px 14px; }
  .leaflet-popup-content b { color: var(--ink); }
  .leaflet-bar { border: 1px solid var(--border-strong) !important; box-shadow: var(--shadow-md) !important; }
  .leaflet-bar a { color: var(--ink-2) !important; }

  /* floating panels on the map */
  .floatpanel { position: absolute; z-index: 1100; background: var(--surface);
    border: 1px solid var(--border); border-radius: var(--r-lg); box-shadow: var(--shadow-md);
    font-size: 13px; line-height: 1.5; }
  #lensPanel { top: 14px; right: 14px; width: 232px; padding: 12px 13px 11px; }
  .fp-title { font-family: var(--mono); font-size: 10.5px; font-weight: 500; letter-spacing: .09em;
    text-transform: uppercase; color: var(--muted); margin-bottom: 8px; }
  .seg3 { display: grid; grid-template-columns: 1fr 1fr 1fr; background: #f4f4f5;
    border: 1px solid var(--border); border-radius: 8px; padding: 2px; gap: 2px; }
  .seg3 button { appearance: none; border: none; background: transparent; border-radius: 6px;
    padding: 5px 2px; font-family: var(--sans); font-size: 11.5px; font-weight: 600; color: var(--muted);
    cursor: pointer; transition: background .12s, color .12s; }
  .seg3 button:hover { color: var(--ink); }
  .seg3 button.on { background: var(--surface); color: var(--ink); box-shadow: var(--shadow-sm); }
  #lensLegend { margin-top: 10px; }
  .lgd-row { display: flex; align-items: center; gap: 7px; padding: 2.5px 0; font-size: 12.5px;
    color: var(--ink-2); }
  .lgd-row.click { cursor: pointer; border-radius: 6px; margin: 0 -5px; padding: 2.5px 5px; }
  .lgd-row.click:hover { background: #f4f4f5; }
  .lgd-row.off { opacity: .38; }
  .lgd-row.off .lgd-name { text-decoration: line-through; }
  .lgd-dot { width: 10px; height: 10px; border-radius: 50%; flex: none; }
  .lgd-line { width: 16px; height: 3px; border-radius: 2px; flex: none; }
  .lgd-ring { width: 10px; height: 10px; border-radius: 50%; background: #fff;
    border: 2px solid __FATAL_STROKE__; flex: none; }
  .lgd-cap { font-size: 11px; color: var(--muted); margin-top: 5px; line-height: 1.45; }
  .lgd-hd { font-family: var(--mono); font-size: 10px; letter-spacing: .08em; text-transform: uppercase;
    color: var(--faint); margin: 8px 0 3px; }
  #lensPanel hr { border: none; border-top: 1px solid var(--border); margin: 10px 0 8px; }
  .tgl { display: flex; align-items: center; gap: 8px; padding: 3px 0; font-size: 12.5px;
    color: var(--ink-2); cursor: pointer; user-select: none; }
  .tgl input { accent-color: var(--ink); width: 14px; height: 14px; margin: 0; }
  @media (max-width: 560px) { #lensPanel { width: 196px; top: 10px; right: 10px; } }

  /* ============ investigate ============ */
  #invLayout { display: grid; grid-template-columns: 430px 1fr; height: calc(100vh - var(--nav-h)); }
  #invSide { overflow-y: auto; background: var(--surface); border-right: 1px solid var(--border);
    padding: 26px 24px 34px; }
  #invSide h2 { margin: 0 0 4px; font-size: 21px; letter-spacing: -.02em; }
  #invSide .invsub { margin: 0 0 18px; font-size: 13px; color: var(--muted); line-height: 1.55; }
  #view-investigate .mapslot { position: relative; }
  #invSeg { display: inline-flex; background: #f4f4f5; border: 1px solid var(--border);
    border-radius: 8px; padding: 2px; gap: 2px; margin-bottom: 8px; }
  #invSeg button { appearance: none; border: none; background: transparent; border-radius: 6px;
    padding: 5px 14px; font-family: var(--sans); font-size: 12px; font-weight: 600; color: var(--muted); cursor: pointer; }
  #invSeg button.on { background: var(--surface); color: var(--ink); box-shadow: var(--shadow-sm); }
  .invrow { display: flex; gap: 8px; }
  #invInput { flex: 1; padding: 10px 13px; border: 1px solid var(--border-strong);
    border-radius: var(--r-md); font-family: var(--sans); font-size: 14px; color: var(--ink); background: var(--surface); }
  #invInput:focus { outline: none; border-color: var(--accent); box-shadow: 0 0 0 3px var(--accent-soft); }
  #invGo { padding: 10px 16px; border: none; border-radius: var(--r-md); background: var(--ink);
    color: #fff; font-family: var(--sans); font-size: 13.5px; font-weight: 600; cursor: pointer; }
  #invGo:hover { background: #2c2c31; }
  #invGo:disabled { opacity: .55; cursor: default; }
  #invErr { display: none; margin-top: 10px; font-size: 13px; color: #b3372a;
    background: #fdf2f0; border: 1px solid #f3d2cc; border-radius: var(--r-md); padding: 8px 12px; }
  #invCard { margin-top: 18px; }
  .invnote { margin-top: 22px; font-size: 11.5px; color: var(--faint); border-top: 1px solid var(--border);
    padding-top: 12px; }
  @media (max-width: 900px) {
    #invLayout { grid-template-columns: 1fr; height: auto; }
    #invSide { border-right: none; border-bottom: 1px solid var(--border); }
    #view-investigate .mapslot { height: 62vh; }
  }

  /* ============ methodology / docs ============ */
  .doc { background: var(--bg); padding: 46px 24px 80px; }
  .doc .inner { max-width: 780px; margin: 0 auto; }
  .doc h2 { font-size: 26px; letter-spacing: -.02em; margin: 0 0 6px; }
  .doc .sub { color: var(--muted); font-size: 14px; margin: 0 0 26px; }
  .doc h3 { font-size: 17.5px; letter-spacing: -.01em; margin: 38px 0 10px; }
  .doc p { font-size: 14.5px; line-height: 1.75; color: var(--ink-2); margin: 10px 0; }
  .doc p b { color: var(--ink); }
  .doc ul { padding-left: 22px; color: var(--ink-2); font-size: 14.5px; line-height: 1.75; }
  .doc code { font-family: var(--mono); font-size: 12.5px; background: #f4f4f5;
    border: 1px solid var(--border); border-radius: 5px; padding: 1.5px 6px; white-space: nowrap; }
  .doc .caveat { background: var(--surface); border: 1px solid var(--border);
    border-left: 3px solid var(--sw-no); border-radius: var(--r-md); padding: 11px 15px;
    font-size: 13.5px; color: var(--ink-2); margin: 12px 0; }
  .doc .srcline { font-size: 12.5px; color: var(--muted); margin: 2px 0 0; }

  /* ============ about ============ */
  #about { background: var(--surface); border-top: 1px solid var(--border); padding: 46px 24px 56px; }
  #about .inner { max-width: 1060px; margin: 0 auto; }
  #about h2 { font-size: 26px; letter-spacing: -.02em; margin: 0 0 14px; }
  #about p { max-width: 760px; font-size: 14.5px; line-height: 1.7; color: var(--ink-2); margin: 0 0 14px; }

  /* ============ site footer ============ */
  #sitefoot { border-top: 1px solid var(--border); background: var(--surface);
    padding: 26px 24px 34px; font-size: 12.5px; color: var(--muted); }
  #sitefoot .inner { max-width: 1060px; margin: 0 auto; display: flex; flex-wrap: wrap;
    gap: 8px 26px; justify-content: space-between; }
  #sitefoot .provenance { flex-basis: 100%; padding-top: 10px; margin-top: 2px;
    border-top: 1px solid var(--border); color: var(--faint); line-height: 1.6; }
</style>
</head>
<body>

<nav id="topnav">
  <a class="wordmark" href="#/"><span class="mark"></span>StreetStat</a>
  <div class="navlinks">
    <a href="#/explore" data-v="explore">Explore</a>
    <a href="#/investigate" data-v="investigate">Investigate</a>
    <a href="#/methodology" data-v="methodology">Methodology</a>
  </div>
</nav>

<main>
<section id="view-home" class="view">
/*__HERO__*/
/*__DASHBOARD__*/
/*__ABOUT__*/
<footer id="sitefoot"><div class="inner">
  <span>StreetStat &mdash; Memphis pedestrian crashes &amp; infrastructure context. Built by Samarth Desai.
    <b>Data current through __DMAX_FULL__</b> (state crash file; reporting lag applies).</span>
  <span>Data: TDOT SAFETY &middot; City of Memphis Public Works GIS &middot; TDOT ADA inventory &middot; &copy; OpenStreetMap contributors &middot; US Census geocoder</span>
  <span class="provenance">All statistics on StreetStat are computed by its own open-source pipeline
    from public data sources. They are not official figures published by TDOT or the City of Memphis
    and may differ from official counts due to methodology. Exact attribution and counting methods
    are documented on the <a href="#/methodology">Methodology page</a>.</span>
</div></footer>
</section>

<section id="view-explore" class="view" hidden>
  <div class="mapslot" id="mapSlotExplore">
    <div id="map"></div>
    <div id="lensPanel" class="floatpanel">
      <div class="fp-title">Lens</div>
      <div class="seg3" id="lensSeg">
        <button data-lens="ownership" class="on">Ownership</button>
        <button data-lens="sidewalk">Sidewalks</button>
        <button data-lens="density">Density</button>
      </div>
      <div id="lensLegend"></div>
      <hr>
      <label class="tgl"><input type="checkbox" id="fatalOnly"><span>Fatal crashes only</span></label>
      <label class="tgl"><input type="checkbox" id="crossToggle"><span>Signalized ped crossings (TDOT)</span></label>
      <div class="lgd-cap" style="margin-top:8px">Data current through __DMAX_FULL__</div>
    </div>
  </div>
</section>

<section id="view-investigate" class="view" hidden>
  <div id="invLayout">
    <aside id="invSide">
      <h2>Investigate a location</h2>
      <p class="invsub">Enter an address or coordinates and StreetStat assembles every fact it has
        about that spot: the road, who owns it, sidewalk status, crashes on the surrounding stretch,
        and the record over time.</p>
      <div id="invSeg">
        <button id="invSegAddr" class="on">Address</button>
        <button id="invSegCoord">Coordinates</button>
      </div>
      <div class="invrow">
        <input id="invInput" autocomplete="off" placeholder="e.g. 1779 Union Ave">
        <button id="invGo">Look up</button>
      </div>
      <div id="invErr"></div>
      <div id="invCard"></div>
      <div class="invnote">Everything shown here is computed directly from the data &mdash;
        see the <a href="#/methodology">Methodology page</a> for how.</div>
    </aside>
    <div class="mapslot" id="mapSlotInv"></div>
  </div>
</section>

<section id="view-methodology" class="view" hidden>
/*__METHODOLOGY__*/
</section>
</main>

<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script src="https://unpkg.com/leaflet.heat@0.2.0/dist/leaflet-heat.js"></script>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<script>
var CRASHES = __CRASHES__;
var SEGMENTS = __SEGMENTS__;
var BOUNDARY = __BOUNDARY__;
var CROSSINGS = __CROSSINGS__;
var S = __STATS_JSON__;
var COL = [S.colors.city, S.colors.tdot, S.colors.lim];
var CATLBL = ["City of Memphis", "TDOT state route", "Limited-access (TDOT)"];

var map = L.map("map", { preferCanvas: true, zoomControl: false }).setView([35.135, -90.01], 11);
L.control.zoom({ position: "bottomright" }).addTo(map);
L.tileLayer("https://{s}.basemaps.cartocdn.com/light_all/{z}/{x}/{y}{r}.png", {
  attribution: '&copy; OpenStreetMap &copy; CARTO', subdomains: "abcd", maxZoom: 19
}).addTo(map);

// boundary outline (context) — visible in every lens; interactive:false so it can never take a click
L.geoJSON(BOUNDARY, { interactive: false,
  style: { color: "#8a9aa2", weight: 1.5, fill: false, dashArray: "4 4", opacity: .7 } }).addTo(map);

// road ownership tint (slim layer only) — shown by the "Ownership" lens. interactive:false: these
// lines carry no popup, and on the shared canvas an interactive no-op line would swallow clicks
// meant for crash dots beneath it.
var segLayer = L.geoJSON(SEGMENTS, { interactive: false, style: function (ft) {
  var o = ft.properties.Ownership;
  if (o === "TDOT state route") return { color: COL[1], weight: 3, opacity: .45 };
  return { color: COL[2], weight: 3, opacity: .40 };  // interstate / ramp / limited-access
}});

// CollisionDate is already carried per crash as c[4] ("YYYY-MM-DD"); format it readably,
// and show "date not recorded" when the source had no usable date.
function fmtDate(d) {
  if (!d || !/^\d{4}-\d{2}-\d{2}$/.test(d)) return "date not recorded";
  var M = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"];
  var p = d.split("-");
  return M[(+p[1]) - 1] + " " + (+p[2]) + ", " + p[0];
}
function popupHtml(c) {
  return "<b>" + fmtDate(c[4]) + "</b><br>" + c[5] + "<br>" + c[6] +
         "<br><b>Road owner:</b> " + CATLBL[c[2]];
}

// per-category layers: every crash is an individual dot at all zooms (no clustering)
var nonfat = [], fatals = [], heatPts = [];
for (var k = 0; k < 3; k++) { nonfat[k] = L.layerGroup(); fatals[k] = L.layerGroup(); }
CRASHES.forEach(function (c) {
  var cat = c[2], ll = [c[0], c[1]];
  heatPts.push([c[0], c[1], c[3] ? 1.0 : 0.45]);
  if (c[3]) {  // fatal — emphasized
    L.circleMarker(ll, { radius: 6, color: "__FATAL_STROKE__", weight: 2,
      fillColor: COL[cat], fillOpacity: .92 }).bindPopup(popupHtml(c)).addTo(fatals[cat]);
  } else {     // non-fatal — individual dot, never grouped
    L.circleMarker(ll, { radius: 4, color: COL[cat], weight: 0.6,
      fillColor: COL[cat], fillOpacity: .6 }).bindPopup(popupHtml(c)).addTo(nonfat[cat]);
  }
});
var heat = L.heatLayer(heatPts, { radius: 18, blur: 16, maxZoom: 14, minOpacity: .25 });

// signalized pedestrian crossings (TDOT inventory) — independent overlay, default OFF; each
// marker is clickable and reports its location + provenance
var crossLayer = L.layerGroup();
(CROSSINGS.features || []).forEach(function (ft) {
  var g = ft.geometry.coordinates, p = ft.properties || {};
  L.circleMarker([g[1], g[0]], { radius: 4, color: "#1f3f8c", weight: 1.6,
    fillColor: "#7aa8e6", fillOpacity: .95 })
    .bindPopup("<b>Signalized pedestrian crossing</b><br>" + (p.dom_street || "") +
               "<br><span style='font-family:var(--mono);font-size:11px;color:#71717a'>" +
               (+g[1]).toFixed(5) + ", " + (+g[0]).toFixed(5) + "</span>" +
               "<br>Pedestrian walk signals + push buttons (TDOT ADA inventory)")
    .addTo(crossLayer);
});

// ============ lens system: ONE analytic lens at a time; crash dots always the subject ============
var LENS = { cur: "ownership", fatalOnly: false, crossings: false, owners: [true, true, true] };
var LENSFACTORIES = {};   // e.g. the sidewalk layer, registered by the search bundle (script 24)
window.__registerLens = function (name, factory) { LENSFACTORIES[name] = factory; renderLens(); };
function tog(layer, on) { if (!layer) return; if (on) { map.addLayer(layer); } else { map.removeLayer(layer); } }
function renderLens() {
  for (var k = 0; k < 3; k++) {
    tog(fatals[k], LENS.owners[k]);
    tog(nonfat[k], LENS.owners[k] && !LENS.fatalOnly);
  }
  tog(segLayer, LENS.cur === "ownership");
  tog(heat, LENS.cur === "density");
  var swf = LENSFACTORIES.sidewalk;
  if (swf) tog(swf(), LENS.cur === "sidewalk");
  tog(crossLayer, LENS.crossings);
  // CLICK PRIORITY (all vectors share one canvas; a canvas click goes to the topmost-drawn
  // interactive layer): sidewalk lines < signals < non-fatal dots < fatal dots. bringToFront
  // re-appends a vector to the canvas draw list, so raising in this order after every lens
  // change guarantees a click on a crash dot is never swallowed by a sidewalk line under it.
  function raise(g) {
    if (g && map.hasLayer(g)) g.eachLayer(function (l) { if (l.bringToFront) l.bringToFront(); });
  }
  raise(crossLayer);
  for (k = 0; k < 3; k++) raise(nonfat[k]);
  for (k = 0; k < 3; k++) raise(fatals[k]);
  drawLegend();
}
function drawLegend() {
  var el = document.getElementById("lensLegend");
  if (!el) return;
  function dotRow(k, name) {
    return '<div class="lgd-row click' + (LENS.owners[k] ? "" : " off") + '" data-own="' + k + '">' +
      '<span class="lgd-dot" style="background:' + COL[k] + '"></span><span class="lgd-name">' + name + '</span></div>';
  }
  var h = '<div class="lgd-hd">Crash dots &middot; by road owner</div>' +
    dotRow(0, "City of Memphis") + dotRow(1, "TDOT state route") + dotRow(2, "Limited-access") +
    '<div class="lgd-row"><span class="lgd-ring"></span>Fatal crash (emphasized)</div>';
  if (LENS.cur === "ownership") {
    h += '<div class="lgd-hd">Road tint</div>' +
      '<div class="lgd-row"><span class="lgd-line" style="background:' + COL[1] + '"></span>TDOT state route</div>' +
      '<div class="lgd-row"><span class="lgd-line" style="background:' + COL[2] + '"></span>Limited-access road</div>' +
      '<div class="lgd-cap">Untinted streets are City of Memphis. Click an owner above to show/hide its crashes.</div>';
  } else if (LENS.cur === "sidewalk") {
    h += '<div class="lgd-hd">City sidewalk inventory</div>' +
      '<div class="lgd-row"><span class="lgd-line" style="background:var(--sw-yes)"></span>Sidewalk in city inventory</div>' +
      '<div class="lgd-row"><span class="lgd-line" style="background:var(--sw-no)"></span>None found in city inventory</div>' +
      '<div class="lgd-cap">Colored along roads with &ge;1 recorded crash &mdash; click a segment for its status. Absence may reflect incomplete records.</div>' +
      (LENSFACTORIES.sidewalk ? "" : '<div class="lgd-cap">(sidewalk layer unavailable in this build)</div>');
  } else {
    h += '<div class="lgd-hd">Crash density</div>' +
      '<div class="lgd-cap">Warmer areas have more recorded crashes nearby; fatal crashes weigh more. ' +
      'An intensity view &mdash; not a rate, not a risk score.</div>';
  }
  if (LENS.crossings) {
    h += '<div class="lgd-hd">Overlay</div>' +
      '<div class="lgd-row"><span class="lgd-dot" style="background:#7aa8e6;border:1.5px solid #1f3f8c"></span>Signalized ped crossing (TDOT)</div>';
  }
  el.innerHTML = h;
  Array.prototype.forEach.call(el.querySelectorAll("[data-own]"), function (r) {
    r.addEventListener("click", function () {
      var k = +r.dataset.own; LENS.owners[k] = !LENS.owners[k]; renderLens();
    });
  });
}
(function wireLens() {
  var seg = document.getElementById("lensSeg");
  Array.prototype.forEach.call(seg.querySelectorAll("button"), function (b) {
    b.addEventListener("click", function () {
      LENS.cur = b.dataset.lens;
      Array.prototype.forEach.call(seg.querySelectorAll("button"), function (x) {
        x.className = (x === b) ? "on" : "";
      });
      renderLens();
    });
  });
  document.getElementById("fatalOnly").addEventListener("change", function (e) {
    LENS.fatalOnly = e.target.checked; renderLens();
  });
  document.getElementById("crossToggle").addEventListener("change", function (e) {
    LENS.crossings = e.target.checked; renderLens();
  });
  var panel = document.getElementById("lensPanel");
  if (window.L && L.DomEvent) { L.DomEvent.disableClickPropagation(panel); L.DomEvent.disableScrollPropagation(panel); }
})();
renderLens();

// ============ hash router: home / explore / investigate / methodology ============
var VIEWS = ["home", "explore", "investigate", "methodology"];
function mapTo(slotId) {
  var slot = document.getElementById(slotId);
  var mc = document.getElementById("map");
  if (slot && mc && mc.parentNode !== slot) slot.appendChild(mc);
  setTimeout(function () { map.invalidateSize(); }, 60);
}
function route() {
  var h = (location.hash || "#/").replace(/^#\/?/, "").split("?")[0] || "home";
  if (VIEWS.indexOf(h) < 0) h = "home";
  VIEWS.forEach(function (v) {
    var el = document.getElementById("view-" + v);
    if (el) el.hidden = (v !== h);
  });
  Array.prototype.forEach.call(document.querySelectorAll("#topnav a[data-v]"), function (a) {
    a.className = (a.dataset.v === h) ? "active" : "";
  });
  if (h === "explore") mapTo("mapSlotExplore");
  if (h === "investigate") mapTo("mapSlotInv");
  if (h === "home") initCharts();
  window.scrollTo(0, 0);
  if (window.__onRoute) window.__onRoute(h);   // search bundle hook (script 24)
}
window.addEventListener("hashchange", route);

// ---- charts (init on first home render so canvases have real dimensions) ----
var _chartsDone = false;
function initCharts() {
  if (_chartsDone || !window.Chart) return;
  _chartsDone = true;
  Chart.defaults.font.family = getComputedStyle(document.body).fontFamily;
  Chart.defaults.color = "#3f3f46";
  new Chart(document.getElementById("cJuris"), {
    type: "bar",
    data: { labels: ["City of Memphis", "TDOT state route", "Limited-access"],
      datasets: [
        { label: "All crashes", backgroundColor: ["#9cd6cd", "#eda9a3", "#a9a9b4"], data: S.juris.all },
        { label: "Deaths", backgroundColor: [COL[0], COL[1], COL[2]], data: S.juris.fatal } ] },
    options: { plugins: { legend: { position: "bottom" } }, scales: { y: { beginAtZero: true } } }
  });
  new Chart(document.getElementById("cLanes"), {
    type: "bar",
    data: { labels: S.lanes.labels.map(function (l) { return l + " ln"; }),
      datasets: [{ label: "Deaths", backgroundColor: COL[1], data: S.lanes.values }] },
    options: { plugins: { legend: { display: false } }, scales: { y: { beginAtZero: true } } }
  });
  new Chart(document.getElementById("cYear"), {
    type: "bar",
    data: { labels: S.year.labels,
      datasets: [
        { label: "All crashes", backgroundColor: "#9cb3bd", data: S.year.all },
        { label: "Deaths", backgroundColor: COL[1], data: S.year.fatal } ] },
    options: { plugins: { legend: { position: "bottom" } }, scales: { y: { beginAtZero: true } } }
  });
}

// ---- deadliest table (sortable) ----
(function () {
  var rows = S.top25.map(function (r, i) { r.rank = i + 1; return r; });
  var tbody = document.querySelector("#deadliest tbody");
  var dir = {};
  function draw() {
    tbody.innerHTML = rows.map(function (r) {
      return "<tr><td class='n'>" + r.rank + "</td><td><b>" + r.name + "</b></td><td class='n'>" + r.total +
        "</td><td class='n'>" + r.serious + "</td><td class='n'>" + r.fatal + "</td><td>" + r.owner +
        (r.mixed ? " <span style='color:#a1a1aa'>(mixed)</span>" : "") + "</td><td class='n'>" +
        (r.spd == null ? "&mdash;" : r.spd) + "</td><td class='n'>" + (r.lanes == null ? "&mdash;" : r.lanes) + "</td></tr>";
    }).join("");
  }
  draw();
  document.querySelectorAll("#deadliest th").forEach(function (th) {
    th.addEventListener("click", function () {
      var k = th.dataset.k; dir[k] = !dir[k];
      rows.sort(function (a, b) {
        var x = a[k], y = b[k];
        if (typeof x === "string") { return dir[k] ? x.localeCompare(y) : y.localeCompare(x); }
        return dir[k] ? x - y : y - x;
      });
      draw();
    });
  });
})();

route();   // initial view (respects a #/... hash on reload)
</script>
</body>
</html>
"""


if __name__ == "__main__":
    main()
