r"""
08_md_to_docx.py
================

Converts a Markdown file into a formatted Microsoft Word (.docx) document so the
statistics reference can be read and edited in Word.

By default it converts:
    data/processed/novel_statistics.md  ->  data/processed/novel_statistics.docx

You can also pass an input and (optional) output path:
    .\.venv\Scripts\python.exe scripts\08_md_to_docx.py <input.md> [output.docx]

Handles the Markdown this project uses: # / ## / ### headings, **bold**,
*italic*, ~~strikethrough~~, `code`, "- " bullet lists, "1." numbered lists,
> blockquotes, | pipe | tables |, and --- horizontal rules.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_IN = PROJECT_ROOT / "data" / "processed" / "novel_statistics.md"

# Matches the inline spans we support, longest/most-specific markers first.
INLINE_RE = re.compile(r"(\*\*.+?\*\*|~~.+?~~|`.+?`|\*.+?\*)")


def add_inline(paragraph, text):
    """Append `text` to a paragraph, honoring **bold**/*italic*/~~strike~~/`code`."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        run = paragraph.add_run()
        if token.startswith("**") and token.endswith("**"):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith("~~") and token.endswith("~~"):
            run.text = token[2:-2]
            run.font.strike = True
        elif token.startswith("`") and token.endswith("`"):
            run.text = token[1:-1]
            run.font.name = "Consolas"
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run.text = token[1:-1]
            run.italic = True
        else:
            run.text = token


def add_horizontal_rule(doc):
    """Add an empty paragraph with a bottom border to act as a divider."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    p_pr.append(borders)


def is_table_line(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    # e.g. |---|:--:|---|
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]  # rows[1] is the separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], cell_text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i < len(cells):
                cells[i].paragraphs[0].text = ""
                add_inline(cells[i].paragraphs[0], cell_text)


def convert(md_path, docx_path):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line.
        if stripped == "":
            i += 1
            continue

        # Horizontal rule.
        if re.fullmatch(r"-{3,}", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables: gather the consecutive pipe lines.
        if is_table_line(line):
            block = []
            while i < len(lines) and is_table_line(lines[i]):
                block.append(lines[i])
                i += 1
            if len(block) >= 2 and is_separator_row(block[1]):
                add_table(doc, block)
            else:  # not a real table; emit as plain paragraphs
                for b in block:
                    add_inline(doc.add_paragraph(), b)
            continue

        # Headings: # -> Title, ## -> Heading 1, ### -> Heading 2.
        heading = re.match(r"(#{1,3})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            style = {1: "Title", 2: "Heading 1", 3: "Heading 2"}[level]
            p = doc.add_paragraph(style=style)
            add_inline(p, heading.group(2))
            i += 1
            continue

        # Blockquote.
        if stripped.startswith(">"):
            text = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, text)
            i += 1
            continue

        # Unordered list item.
        if re.match(r"[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            i += 1
            continue

        # Ordered list item.
        if re.match(r"\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text)
            i += 1
            continue

        # Plain paragraph.
        add_inline(doc.add_paragraph(), stripped)
        i += 1

    doc.save(docx_path)
    return docx_path


def main():
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_IN
    if len(sys.argv) > 2:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix(".docx")
    out = convert(md_path, docx_path)
    print(f"Converted: {md_path.name}  ->  {out}")


if __name__ == "__main__":
    main()
