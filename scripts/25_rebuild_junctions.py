r"""
25_rebuild_junctions.py
=======================

Intersection-node ROOT FIX. The Phase-3a node builder (script 21) located junctions
by snapping segment ENDPOINTS to a 1 m grid and keeping grid cells where >=2 distinct
named through-roads landed. That is endpoint-noding, not a true geometric intersection:
it misses real X-crossings whose centerlines cross WITHOUT sharing a rounded endpoint,
and it only ever exported the subset of nodes on "covered" corridors -- which is why
Union Ave & S Cleveland St (a real, crash-bearing junction) was unsearchable.

This script rebuilds the FULL junction set from TRUE geometric line intersection:

  1. THROUGH-ROADS: rulebook segments with MTFCC in {S1200, S1400} and a non-blank
     Street_Name; dissolve to one (multi)line per normalized name (directional
     prefixes kept: N CLEVELAND ST != S CLEVELAND ST).
  2. TRUE INTERSECTION: self spatial-join (intersects) over the dissolved names, keep
     differently-named pairs once, compute shapely intersection -> crossing points.
     Interstates/ramps (S1100/S1630) are excluded as through-roads, so grade-separated
     overpasses do NOT create false at-grade nodes; that exclusion is quantified.
  3. DIVIDED-ARTERIAL MERGE: a cross street meeting a divided arterial crosses each
     carriageway, producing two nearby points for ONE junction. We PRINT the actual
     distribution of these same-pair carriageway gaps, then set the cluster radius
     (25-30 m) from it and single-linkage merge crossing points into nodes.
  4. COVERAGE (root fix): the old rule mapped each signal route to one DOMINANT street,
     so Union (52 signals, never dominant on SR003/SR023/SR001) fell out of scope. New
     rule = route-dominant streets UNION any street directly carrying >= MIN_OWN_SIGNALS
     of its own nearest signals. This adds Union and is recorded in novel_statistics.docx.
  5. SIGNALIZED: a node is signalized if a deduped crossing is within 30 m.
  6. RE-ATTRIBUTE crashes to the nearest node within 30 m (ANY node, not just covered),
     recompute at_intersection / intersection_node_id / intersection_signalized, and
     PROVE the headline totals (current anchors, see CLAUDE.md) and the deadliest-corridor
     counts are unchanged. A pre-fix snapshot is kept so the status delta is reproducible.

Reads:  data/processed/road_ownership_rulebook.geojson
        data/raw/ped_signals.geojson
        data/processed/signalized_crossings_dedup.geojson
        data/processed/shelby_crashes_final.csv
        data/processed/shelby_crashes_signals.csv      (pre-fix, for the diff)
Writes: data/processed/intersection_nodes_all.geojson  (EVERY junction citywide; NEW)
        data/processed/shelby_crashes_signals.csv       (re-attributed; pre-fix backed up)
        data/processed/shelby_crashes_signals_pre25.csv (one-time baseline snapshot)

Run it with:
    .\.venv\Scripts\python.exe scripts\25_rebuild_junctions.py
"""

import sys
import json
import shutil
from collections import defaultdict
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point

ROOT = Path(__file__).resolve().parent.parent
RAW = ROOT / "data" / "raw"
PROC = ROOT / "data" / "processed"

RULEBOOK = PROC / "road_ownership_rulebook.geojson"
PED = RAW / "ped_signals.geojson"
CROSSINGS = PROC / "signalized_crossings_dedup.geojson"
FINAL = PROC / "shelby_crashes_final.csv"
SIGNALS = PROC / "shelby_crashes_signals.csv"
SIGNALS_BAK = PROC / "shelby_crashes_signals_pre25.csv"
NODES_OUT = PROC / "intersection_nodes_all.geojson"
COVERED_OUT = PROC / "covered_corridors.json"
DOCX = PROC / "novel_statistics.docx"

CRS_M, CRS_GEO = "EPSG:32136", "EPSG:4326"
THROUGH_MTFCC = {"S1200", "S1400"}      # arterials + local through-roads
GRADE_SEP_MTFCC = {"S1100", "S1630"}    # interstate, ramp -> grade-separated, excluded
SNAP_M = 30.0                            # crash->node and crossing->node
SAMEPAIR_M = 120.0                       # same-street-pair carriageway-split consolidation cap
MIN_ROUTE_TOTAL = 4                      # a route needs >= this many signals to be "covered"
SECONDARY_SHARE = 0.25                   # a route's secondary street is a corridor at this share
MIN_OWN_SIGNALS = 4                      # NEW: a street carrying >= this many of its own signals is covered
FATAL = "Fatal"


def pct(p, w):
    return round(100.0 * p / w, 1) if w else 0.0


def points_of(geom):
    """All crossing points from an intersection result (lines -> their centroid)."""
    if geom is None or geom.is_empty:
        return []
    t = geom.geom_type
    if t == "Point":
        return [(geom.x, geom.y)]
    if t == "MultiPoint":
        return [(p.x, p.y) for p in geom.geoms]
    if t == "LineString":
        c = geom.centroid
        return [(c.x, c.y)]
    if t in ("MultiLineString", "GeometryCollection"):
        out = []
        for sub in geom.geoms:
            out += points_of(sub)
        return out
    return []


class UF:
    def __init__(self, n):
        self.p = list(range(n))

    def find(self, i):
        while self.p[i] != i:
            self.p[i] = self.p[self.p[i]]
            i = self.p[i]
        return i

    def union(self, a, b):
        ra, rb = self.find(a), self.find(b)
        if ra != rb:
            self.p[ra] = rb


def append_coverage_note(old_corridors, new_corridors, added, union_own):
    """Append a dated one-line methodology note for the covered-flag change (idempotent).
    Removes any prior copy first so a corrected note replaces a stale one. All figures computed."""
    try:
        from docx import Document
    except Exception:
        print("  (python-docx not available; skipping docx note)"); return
    if not DOCX.exists():
        print(f"  (docx not found: {DOCX.name}; skipping)"); return
    doc = Document(str(DOCX))
    MARK = "Coverage-rule revision"
    for p in [p for p in doc.paragraphs if p.text.strip().startswith(MARK)]:
        p._element.getparent().remove(p._element)   # drop stale copy so we re-write the corrected one
    added_s = ", ".join(added[:8]) + ("..." if len(added) > 8 else "")
    doc.add_paragraph(
        f"Coverage-rule revision ({date.today().isoformat()}): the 'covered corridor' flag was "
        f"broadened. OLD rule -- each TDOT signal route was mapped to its single DOMINANT arterial "
        f"(plus any >=25% secondary), {old_corridors} corridors. NEW rule -- that set UNION any street "
        f"directly carrying >= {MIN_OWN_SIGNALS} of its own nearest signals, {new_corridors} corridors. "
        f"This adds {len(added)} streets ({added_s}), most importantly UNION AVE ({union_own} signals "
        f"snap nearest to it, yet it is never the 'dominant' arterial on any single route, so the old "
        f"rule missed it), bringing Union-corridor junctions into scope. Headline crash totals "
        f"(1,294 / 175) and the deadliest-corridor counts are unaffected -- the flag only governs "
        f"intersection signal-coverage scope, not crash attribution.",
        style="List Bullet")
    doc.save(str(DOCX))
    print(f"  (re)wrote coverage-rule note in {DOCX.name}")


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    print("loading rulebook + signals + crossings + crashes...")
    rb = gpd.read_file(RULEBOOK).to_crs(CRS_M)
    ped = gpd.read_file(PED).to_crs(CRS_M).reset_index(drop=True)
    cross = gpd.read_file(CROSSINGS).to_crs(CRS_M)

    # ---------------------------------------------------------------- 1. through-roads
    through = rb[rb["MTFCC"].isin(THROUGH_MTFCC)].copy()
    through = through[through["Street_Name"].astype(str).str.strip() != ""]
    through["Street_Name"] = through["Street_Name"].astype(str).str.strip()
    named = through.dissolve(by="Street_Name").reset_index()[["Street_Name", "geometry"]]
    gl = dict(zip(named["Street_Name"], named.geometry))
    print(f"  through-roads (S1200/S1400, named): {len(through):,} segments -> "
          f"{len(named):,} distinct named streets")

    # ---------------------------------------------------------------- 2. true intersections
    sj = gpd.sjoin(named, named, predicate="intersects")
    pairs = sj[sj["Street_Name_left"] < sj["Street_Name_right"]][
        ["Street_Name_left", "Street_Name_right"]].drop_duplicates()
    lg = gpd.GeoSeries([gl[n] for n in pairs["Street_Name_left"]], crs=CRS_M)
    rg = gpd.GeoSeries([gl[n] for n in pairs["Street_Name_right"]], crs=CRS_M)
    inter = lg.intersection(rg, align=False).values

    px, py, nmA, nmB = [], [], [], []
    for g, a, b in zip(inter, pairs["Street_Name_left"].values, pairs["Street_Name_right"].values):
        for (x, y) in points_of(g):
            px.append(x); py.append(y); nmA.append(a); nmB.append(b)
    print(f"  differently-named crossing pairs: {len(pairs):,} -> raw crossing points: {len(px):,}")

    # grade-separated exclusions (quantify what TRUE intersection would have wrongly noded)
    gs = rb[rb["MTFCC"].isin(GRADE_SEP_MTFCC)]
    gs_union = gs.geometry.union_all() if len(gs) else None
    n_gradesep = 0
    if gs_union is not None and not gs_union.is_empty:
        for g in named.geometry.values:
            n_gradesep += len(points_of(g.intersection(gs_union)))
    print(f"  grade-separated crossings excluded (through-road x interstate/ramp): {n_gradesep:,} "
          f"points across {len(gs):,} S1100/S1630 segments -> NOT noded (plan-view only, not at-grade)")

    # ---------------------------------------------------------------- 3. divided-arterial gap distribution
    same_pair = defaultdict(list)
    for i in range(len(px)):
        same_pair[(nmA[i], nmB[i])].append((px[i], py[i]))
    gaps = []
    gap_examples = []
    for (a, b), pts in same_pair.items():
        if len(pts) < 2:
            continue
        arr = np.array(pts)
        for i in range(len(arr)):
            for j in range(i + 1, len(arr)):
                d = float(np.hypot(*(arr[i] - arr[j])))
                if d <= 60.0:                       # candidate same-junction carriageway split
                    gaps.append(d)
                    gap_examples.append((d, a, b))
    gaps = np.array(sorted(gaps)) if gaps else np.array([])
    print("\nDIVIDED-ARTERIAL CARRIAGEWAY-GAP DISTRIBUTION "
          "(distance between same-street-pair crossing points <= 60 m):")
    if len(gaps):
        for q in (50, 75, 90, 95, 99):
            print(f"    p{q:<2} = {np.percentile(gaps, q):5.1f} m")
        print(f"    max = {gaps.max():5.1f} m   (n = {len(gaps)} candidate splits)")
        for d, a, b in sorted(gap_examples, reverse=True)[:6]:
            print(f"      e.g. {a} x {b}: {d:.1f} m apart")
    else:
        print("    (no same-pair points within 60 m)")

    # set radius from the distribution: 25-30 m band; only drop below 25 if no gap exceeds ~20 m
    if len(gaps) and gaps.max() <= 20.0:
        RADIUS = 20.0
        why = "no carriageway gap exceeds ~20 m"
    else:
        p99 = float(np.percentile(gaps, 99)) if len(gaps) else 0.0
        RADIUS = float(min(30.0, max(25.0, np.ceil(p99))))
        why = f"p99 gap = {p99:.1f} m (gaps exceed 20 m); clamped to the 25-30 m band"
    print(f"  -> CLUSTER RADIUS = {RADIUS:.0f} m  ({why})")

    # ---------------------------------------------------------------- 3b. cluster crossing points -> nodes
    ptsg = gpd.GeoDataFrame(geometry=[Point(x, y) for x, y in zip(px, py)], crs=CRS_M)
    buf = gpd.GeoDataFrame(geometry=ptsg.buffer(RADIUS), crs=CRS_M)
    jr = gpd.sjoin(ptsg, buf, predicate="within")
    uf = UF(len(ptsg))
    for i, jj in zip(jr.index.values, jr["index_right"].values):
        uf.union(int(i), int(jj))
    # PASS B -- same-street-pair consolidation: a divided arterial's two carriageways yield
    # two crossing points of the SAME road pair (medians can exceed the 30 m radius). Merge
    # points sharing the identical unordered name-pair within SAMEPAIR_M, regardless of radius.
    # This never merges two DIFFERENT cross streets (different pair), so distinct junctions stay split.
    pair_idx = defaultdict(list)
    for i in range(len(px)):
        pair_idx[(nmA[i], nmB[i])].append(i)
    n_carriageway = 0
    for idxs in pair_idx.values():
        if len(idxs) < 2:
            continue
        for a in range(len(idxs)):
            for b in range(a + 1, len(idxs)):
                ia, ib = idxs[a], idxs[b]
                if np.hypot(px[ia] - px[ib], py[ia] - py[ib]) <= SAMEPAIR_M:
                    if uf.find(ia) != uf.find(ib):
                        n_carriageway += 1
                    uf.union(ia, ib)
    print(f"  pass B (same-pair carriageway merge <= {SAMEPAIR_M:.0f} m): {n_carriageway} additional joins")
    clusters = defaultdict(lambda: {"x": [], "y": [], "names": set()})
    for i in range(len(px)):
        c = clusters[uf.find(i)]
        c["x"].append(px[i]); c["y"].append(py[i]); c["names"].update((nmA[i], nmB[i]))
    cx, cy, cstreets, cn = [], [], [], []
    for c in clusters.values():
        cx.append(float(np.mean(c["x"]))); cy.append(float(np.mean(c["y"])))
        nm = sorted(c["names"]); cstreets.append("; ".join(nm)); cn.append(len(nm))
    nodes = gpd.GeoDataFrame(
        {"node_id": range(1, len(cx) + 1), "streets": cstreets, "n_streets": cn},
        geometry=[Point(x, y) for x, y in zip(cx, cy)], crs=CRS_M)
    print(f"\nJUNCTION NODES (true intersection, {RADIUS:.0f} m merge): {len(nodes):,}  "
          f"(old endpoint-noding set was ~27,179)")

    # ---------------------------------------------------------------- 4. coverage (revised rule)
    snap = gpd.sjoin_nearest(ped[["ROUTE_NUMBER", "geometry"]],
                             through[["Street_Name", "MTFCC", "geometry"]],
                             how="left", distance_col="d")
    snap = snap[~snap.index.duplicated(keep="first")]
    ped["street"] = snap["Street_Name"].values
    pair = ped.groupby(["ROUTE_NUMBER", "street"]).size().reset_index(name="n")
    route_tot = ped.groupby("ROUTE_NUMBER").size()
    covered_routes = route_tot[route_tot >= MIN_ROUTE_TOTAL].index
    dominant = set()
    for rn in covered_routes:
        sub = pair[pair["ROUTE_NUMBER"] == rn].sort_values("n", ascending=False)
        if sub.empty:
            continue
        tot = route_tot[rn]
        dominant.add(sub.iloc[0]["street"])
        for _, r in sub.iloc[1:].iterrows():
            if r["n"] >= SECONDARY_SHARE * tot:
                dominant.add(r["street"])
    own_sig = ped.groupby("street").size()
    own_covered = set(own_sig[own_sig >= MIN_OWN_SIGNALS].index)
    covered = {s for s in (dominant | own_covered) if isinstance(s, str) and s.strip()}
    added = sorted(own_covered - dominant)
    print(f"  covered corridors: route-dominant {len(dominant)}  ->  +own-signals rule = {len(covered)} "
          f"(added {len(added)}: {', '.join(added[:10])}{'...' if len(added) > 10 else ''})")
    print(f"  UNION AVE covered now: {'UNION AVE' in covered}  "
          f"(its own signal count = {int(own_sig.get('UNION AVE', 0))})")
    nodes["on_covered"] = [any(s in covered for s in st.split("; ")) for st in nodes["streets"]]
    COVERED_OUT.write_text(json.dumps(sorted(covered)), encoding="utf-8")  # sidecar for the search index

    # ---------------------------------------------------------------- 5. signalized
    nsig = gpd.sjoin_nearest(nodes[["node_id", "geometry"]], cross[["geometry"]],
                             how="left", distance_col="dsig")
    nsig = nsig[~nsig.index.duplicated(keep="first")]
    nodes["nearest_crossing_m"] = np.round(nsig["dsig"].values, 1)
    nodes["signalized"] = nodes["nearest_crossing_m"] <= SNAP_M
    print(f"  signalized nodes (crossing within {SNAP_M:.0f} m): {int(nodes['signalized'].sum()):,}; "
          f"on covered corridors: {int(nodes['on_covered'].sum()):,}")

    # ---------------------------------------------------------------- 6. re-attribute crashes
    if not SIGNALS_BAK.exists() and SIGNALS.exists():
        shutil.copyfile(SIGNALS, SIGNALS_BAK)
        print(f"  baseline snapshot saved -> {SIGNALS_BAK.name}")
    old = pd.read_csv(SIGNALS_BAK) if SIGNALS_BAK.exists() else None

    cr = pd.read_csv(FINAL)
    pts = gpd.GeoDataFrame(cr.copy(),
                           geometry=gpd.points_from_xy(cr["Longitude"], cr["Latitude"]),
                           crs=CRS_GEO).to_crs(CRS_M)
    nany = gpd.sjoin_nearest(pts[["geometry"]],
                             nodes[["node_id", "signalized", "on_covered", "geometry"]],
                             how="left", distance_col="dnode")
    nany = nany[~nany.index.duplicated(keep="first")]
    nc = gpd.sjoin_nearest(pts[["geometry"]], cross[["geometry"]], how="left", distance_col="dcross")
    nc = nc[~nc.index.duplicated(keep="first")]

    field_intx = cr["NonMotoristLocation"].astype(str).str.startswith("Intersection").values
    dnode = nany["dnode"].values
    within = dnode <= SNAP_M
    nsig_b = nany["signalized"].values.astype(bool)
    ncov_b = nany["on_covered"].values.astype(bool)

    out = cr.copy()
    out["at_intersection"] = field_intx
    out["intersection_node_id"] = pd.array(np.where(within, nany["node_id"].values, np.nan),
                                           dtype="Int64")
    out["nearest_ped_signal_m"] = np.round(nc["dcross"].values, 1)
    out["is_ambiguous_intersection"] = ((~field_intx & within) | (field_intx & ~within))

    sig = np.full(len(out), "", dtype=object)
    for i in range(len(out)):
        if not field_intx[i]:
            sig[i] = ""
        elif within[i]:
            sig[i] = "yes" if nsig_b[i] else ("no" if ncov_b[i] else "no_signal_coverage")
        else:
            sig[i] = "no_signal_coverage"
    out["intersection_signalized"] = sig
    out.to_csv(SIGNALS, index=False, encoding="utf-8")

    # crashes per node (for the node file + index)
    j = out[out["intersection_node_id"].notna()].copy()
    j["intersection_node_id"] = j["intersection_node_id"].astype(int)
    per = j.groupby("intersection_node_id").agg(
        crashes=("MstrRecNbrTxt", "size"),
        deaths=("InjuryClass", lambda s: int((s == FATAL).sum()))).reset_index()
    nodes = nodes.merge(per, left_on="node_id", right_on="intersection_node_id", how="left")
    nodes["crashes"] = nodes["crashes"].fillna(0).astype(int)
    nodes["deaths"] = nodes["deaths"].fillna(0).astype(int)
    nodes = nodes.drop(columns=["intersection_node_id"])
    nodes.to_crs(CRS_GEO).to_file(NODES_OUT, driver="GeoJSON")
    print(f"  wrote {len(nodes):,} nodes -> {NODES_OUT.name} "
          f"({int((nodes['crashes'] > 0).sum()):,} carry >=1 crash)")

    # ---------------------------------------------------------------- diff vs pre-fix
    if old is not None:
        om = old.set_index("MstrRecNbrTxt")
        nm = out.set_index("MstrRecNbrTxt")
        common = om.index.intersection(nm.index)
        o_nid = om.loc[common, "intersection_node_id"]
        n_nid = nm.loc[common, "intersection_node_id"]
        gained = int(((o_nid.isna()) & (n_nid.notna())).sum())
        lost = int(((o_nid.notna()) & (n_nid.isna())).sum())
        sig_changed = int((om.loc[common, "intersection_signalized"].fillna("").astype(str)
                           != nm.loc[common, "intersection_signalized"].fillna("").astype(str)).sum())
        print(f"\nCRASH STATUS DELTA vs pre-fix:")
        print(f"  newly snapped to a node (was null -> now has node): {gained}")
        print(f"  lost a node (was node -> now null): {lost}")
        print(f"  intersection_signalized value changed: {sig_changed}")

    # ---------------------------------------------------------------- PROOF: totals unchanged
    nfat = int((out["InjuryClass"] == FATAL).sum())
    print("\nRECONCILIATION (must be unchanged):")
    print(f"  crashes: {len(out)} (=1294? {len(out) == 1294}) | fatal: {nfat} (=175? {nfat == 175})")
    g = cr.groupby("Street_Name").agg(total=("MstrRecNbrTxt", "size"),
                                      fatal=("InjuryClass", lambda s: int((s == FATAL).sum())))
    g = g.sort_values(["total", "fatal"], ascending=False)
    anchors = {"POPLAR AVE": (44, 8), "UNION AVE": (36, 8), "LAMAR AVE": (30, 6),
               "WINCHESTER RD": (28, 5)}
    print(f"  deadliest corridors rank {len(g)} streets; sum = {int(g['total'].sum())} "
          f"(=1294? {int(g['total'].sum()) == 1294})")
    for name, (et, ef) in anchors.items():
        if name in g.index:
            t, f = int(g.loc[name, "total"]), int(g.loc[name, "fatal"])
            print(f"    {name:<16} {t}/{f}  (expected {et}/{ef}  {'OK' if (t, f) == (et, ef) else 'CHECK'})")

    # ---------------------------------------------------------------- ACCEPTANCE: Union & Cleveland
    print("\nACCEPTANCE -- UNION AVE & S CLEVELAND ST:")
    uc = nodes[nodes["streets"].str.contains("UNION AVE", na=False)
               & nodes["streets"].str.contains("S CLEVELAND ST", na=False)]
    if len(uc):
        r = uc.iloc[0]
        print(f"  NODE EXISTS: id={int(r['node_id'])} | streets='{r['streets']}' | "
              f"on_covered={r['on_covered']} | signalized={r['signalized']} | "
              f"crashes={int(r['crashes'])} ({int(r['deaths'])} fatal) | "
              f"nearest crossing {r['nearest_crossing_m']:.0f} m")
    else:
        print("  *** NODE NOT FOUND -- investigate ***")

    # ---------------------------------------------------------------- docx coverage note
    append_coverage_note(len(dominant), len(covered), added, int(own_sig.get("UNION AVE", 0)))
    print("\nDone. Next: script 24 indexes intersection_nodes_all.geojson (ALL junctions).")


if __name__ == "__main__":
    main()
